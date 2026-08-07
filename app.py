import re
import json
import base64
import datetime
import unicodedata
import urllib.request
import urllib.error
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from io import BytesIO
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill, Border, Side, Font, Alignment
from openpyxl.utils import get_column_letter

try:
    import jpholiday
    HAS_JPHOLIDAY = True
except ImportError:
    HAS_JPHOLIDAY = False

st.set_page_config(
    page_title="献立チェックシステム",
    layout="wide",
    page_icon="🍱"
)

BASE_DIR = Path(__file__).parent
RULES_FILE = BASE_DIR / "rules.txt"
RULES_JSON_FILE = BASE_DIR / "rules.json"
COLOR_RULES_JSON_FILE = BASE_DIR / "color_rules.json"
REPLACE_RULES_JSON_FILE = BASE_DIR / "replace_rules.json"

GITHUB_OWNER = "fukiyamiyazaki-blip"
GITHUB_REPO = "-"
GITHUB_BRANCH = "main"
GITHUB_RULES_PATH = "rules.txt"
GITHUB_RULES_JSON_PATH = "rules.json"
GITHUB_COLOR_RULES_JSON_PATH = "color_rules.json"
GITHUB_REPLACE_RULES_JSON_PATH = "replace_rules.json"

# 色付けルールが1件も登録されていない場合に使う初期セット（従来の固定配色と同一）
DEFAULT_COLOR_GROUPS = [
    {"keywords": ["コーン", "人参", "黄パプリカ", "赤パプリカ", "かぼちゃ"], "color": "FFFF99"},
    {"keywords": ["ほうれん草", "小松菜", "チンゲン菜", "グリンピース",
                  "いんげん", "えだまめ", "ブロッコリー", "ピーマン"],      "color": "CCFFCC"},
    {"keywords": ["木綿豆腐", "焼き豆腐", "油揚げ", "厚揚げ", "大豆"],      "color": "FFD9AD"},
    {"keywords": ["チーズ"],                                              "color": "E6CCFF"},
    {"keywords": ["ちくわ", "かにかま", "ツナ", "赤かまぼこ"],             "color": "CCFFFF"},
    {"keywords": ["ロースハム", "ベーコン", "ウインナー"],                 "color": "FFB3C6"},
]

# Streamlit の C キーショートカット（キャッシュクリア）を無効化
components.html("""
<script>
try {
    function blockCKey(e) {
        if (e.key !== 'c' && e.key !== 'C') return;
        var el = window.parent.document.activeElement;
        if (el) {
            var tag = el.tagName.toUpperCase();
            if (tag === 'INPUT' || tag === 'TEXTAREA' || el.isContentEditable) return;
        }
        e.stopImmediatePropagation();
        e.stopPropagation();
    }
    window.parent.document.addEventListener('keydown', blockCKey, true);
    window.parent.document.addEventListener('keyup', blockCKey, true);
    window.parent.addEventListener('keydown', blockCKey, true);
    window.parent.addEventListener('keyup', blockCKey, true);
} catch(err) {}
</script>
""", height=0, scrolling=False)


def load_rules():
    if RULES_FILE.exists():
        return RULES_FILE.read_text(encoding="utf-8")
    return ""


def save_rules(text):
    RULES_FILE.write_text(text, encoding="utf-8")


def load_rules_list():
    """複数ルールをリストで返す。GitHubを優先し、失敗時のみローカル→rules.txtにフォールバック。"""
    token = get_github_token()
    if token:
        api_url = (
            f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
            f"/contents/{GITHUB_RULES_JSON_PATH}?ref={GITHUB_BRANCH}"
        )
        headers = {
            "Authorization": f"token {token}",
            "Accept": "application/vnd.github.v3+json",
        }
        req = urllib.request.Request(api_url, headers=headers)
        try:
            with urllib.request.urlopen(req) as resp:
                data = json.loads(resp.read())
                content = base64.b64decode(data["content"]).decode("utf-8")
                return json.loads(content).get("rules", [])
        except Exception:
            pass  # GitHub取得失敗 → ローカルにフォールバック

    # ローカルファイルにフォールバック
    if RULES_JSON_FILE.exists():
        try:
            data = json.loads(RULES_JSON_FILE.read_text(encoding="utf-8"))
            return data.get("rules", [])
        except Exception:
            pass
    # 後方互換: rules.txt があれば移行して初期ルールとして返す
    legacy = load_rules()
    if legacy.strip():
        return [{"id": "default", "name": "共通ルール", "text": legacy}]
    return []


def save_rules_list(rules_list):
    """複数ルールをJSON保存（ローカル）。失敗してもエラーを出さない。"""
    try:
        RULES_JSON_FILE.write_text(
            json.dumps({"rules": rules_list}, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
    except Exception:
        pass  # Streamlit Cloud では書き込み不可の場合もあるが GitHub が主ストレージ


def load_color_rules_list():
    """色付けルール（複数セット）をリストで返す。GitHubを優先し、失敗時のみローカルにフォールバック。
    1件も登録がない場合は従来の固定配色をデフォルトルールとして返す（未保存・表示専用）。
    各要素: {"id": str, "name": str, "groups": [{"keywords": [str,...], "color": "RRGGBB"}, ...]}
    """
    token = get_github_token()
    if token:
        api_url = (
            f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
            f"/contents/{GITHUB_COLOR_RULES_JSON_PATH}?ref={GITHUB_BRANCH}"
        )
        headers = {
            "Authorization": f"token {token}",
            "Accept": "application/vnd.github.v3+json",
        }
        req = urllib.request.Request(api_url, headers=headers)
        try:
            with urllib.request.urlopen(req) as resp:
                data = json.loads(resp.read())
                content = base64.b64decode(data["content"]).decode("utf-8")
                rules = json.loads(content).get("color_rules", [])
                if rules:
                    return rules
        except Exception:
            pass  # GitHub取得失敗 → ローカルにフォールバック

    if COLOR_RULES_JSON_FILE.exists():
        try:
            data = json.loads(COLOR_RULES_JSON_FILE.read_text(encoding="utf-8"))
            rules = data.get("color_rules", [])
            if rules:
                return rules
        except Exception:
            pass

    return [{"id": "default", "name": "デフォルト（従来の配色）", "groups": DEFAULT_COLOR_GROUPS}]


def save_color_rules_list(color_rules_list):
    """色付けルールをJSON保存（ローカル）。失敗してもエラーを出さない。"""
    try:
        COLOR_RULES_JSON_FILE.write_text(
            json.dumps({"color_rules": color_rules_list}, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
    except Exception:
        pass


def push_color_rules_list_to_github(color_rules_list):
    """色付けルール（color_rules.json）をGitHubに保存。"""
    token = get_github_token()
    if not token:
        return False, "GitHubトークンが未設定です"

    api_url = (
        f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
        f"/contents/{GITHUB_COLOR_RULES_JSON_PATH}"
    )
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json",
    }

    sha = None
    req = urllib.request.Request(
        api_url + f"?ref={GITHUB_BRANCH}", headers=headers
    )
    try:
        with urllib.request.urlopen(req) as resp:
            sha = json.loads(resp.read())["sha"]
    except Exception:
        pass  # 新規ファイルの場合はSHAなし

    content_text = json.dumps({"color_rules": color_rules_list}, ensure_ascii=False, indent=2)
    body = {"message": "色付けルール管理から更新", "branch": GITHUB_BRANCH,
            "content": base64.b64encode(content_text.encode("utf-8")).decode("utf-8")}
    if sha:
        body["sha"] = sha

    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(api_url, data=payload, headers=headers, method="PUT")
    try:
        with urllib.request.urlopen(req):
            return True, "保存しました（GitHub反映済み）"
    except Exception as e:
        return False, f"GitHub更新エラー: {e}"


def load_replace_rules_list():
    """置換ルール（複数セット）をリストで返す。GitHubを優先し、失敗時のみローカルにフォールバック。
    各要素: {"id": str, "name": str, "pairs": [{"from": str, "to": str}, ...]}
    """
    token = get_github_token()
    if token:
        api_url = (
            f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
            f"/contents/{GITHUB_REPLACE_RULES_JSON_PATH}?ref={GITHUB_BRANCH}"
        )
        headers = {
            "Authorization": f"token {token}",
            "Accept": "application/vnd.github.v3+json",
        }
        req = urllib.request.Request(api_url, headers=headers)
        try:
            with urllib.request.urlopen(req) as resp:
                data = json.loads(resp.read())
                content = base64.b64decode(data["content"]).decode("utf-8")
                return json.loads(content).get("replace_rules", [])
        except Exception:
            pass  # GitHub取得失敗 → ローカルにフォールバック

    if REPLACE_RULES_JSON_FILE.exists():
        try:
            data = json.loads(REPLACE_RULES_JSON_FILE.read_text(encoding="utf-8"))
            return data.get("replace_rules", [])
        except Exception:
            pass
    return []


def save_replace_rules_list(replace_rules_list):
    """置換ルールをJSON保存（ローカル）。失敗してもエラーを出さない。"""
    try:
        REPLACE_RULES_JSON_FILE.write_text(
            json.dumps({"replace_rules": replace_rules_list}, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
    except Exception:
        pass


def push_replace_rules_list_to_github(replace_rules_list):
    """置換ルール（replace_rules.json）をGitHubに保存。"""
    token = get_github_token()
    if not token:
        return False, "GitHubトークンが未設定です"

    api_url = (
        f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
        f"/contents/{GITHUB_REPLACE_RULES_JSON_PATH}"
    )
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json",
    }

    sha = None
    req = urllib.request.Request(
        api_url + f"?ref={GITHUB_BRANCH}", headers=headers
    )
    try:
        with urllib.request.urlopen(req) as resp:
            sha = json.loads(resp.read())["sha"]
    except Exception:
        pass  # 新規ファイルの場合はSHAなし

    content_text = json.dumps({"replace_rules": replace_rules_list}, ensure_ascii=False, indent=2)
    body = {"message": "置換ルール管理から更新", "branch": GITHUB_BRANCH,
            "content": base64.b64encode(content_text.encode("utf-8")).decode("utf-8")}
    if sha:
        body["sha"] = sha

    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(api_url, data=payload, headers=headers, method="PUT")
    try:
        with urllib.request.urlopen(req):
            return True, "保存しました（GitHub反映済み）"
    except Exception as e:
        return False, f"GitHub更新エラー: {e}"


# ─────────────────────────────────────────────
# Python事前計算用ヘルパー
# ─────────────────────────────────────────────

def _parse_date(ds, year):
    m = re.match(r'(\d+)/(\d+)\(', ds)
    if not m:
        return None
    try:
        return datetime.date(year, int(m.group(1)), int(m.group(2)))
    except ValueError:
        return None


def _dow(ds):
    """月=0 … 日=6"""
    m = re.search(r'\(([月火水木金土日])\)', ds)
    return "月火水木金土日".index(m.group(1)) if m else -1


def _is_holiday(d):
    if not HAS_JPHOLIDAY or d is None:
        return False
    try:
        return bool(jpholiday.is_holiday(d))
    except Exception:
        return False


def _split_ing(ing_text):
    """材料テキストを個別トークンに分割"""
    parts = re.split(r'[,、]', ing_text)
    result = []
    _num = re.compile(r'^\d+\.?\d*$')  # 数値のみのトークン（数量・シリアル値）を除外
    for part in parts:
        # 半角・全角カッコも区切り文字として扱う
        for token in re.split(r'[\s　・／/()（）]+', part.strip()):
            t = token.strip()
            if t and t not in ('nan', '') and not _num.match(t):
                result.append(t)
    return result


def _parse_structured(excel_text):
    """
    excel_text を日付ごとの構造体に変換。
    Returns: (year, month_num, sorted_dates, entries)
      entries[date_str] = {'lunch': str, 'snack': str, 'ingredients': str,
                            'snack_am': str, 'snack_pm': str}
      snack_am/snack_pm は「午前おやつ:」「午後おやつ:」が出力される形式（くにみ子ども園等）
      でのみ入る。他形式では空文字のまま（既存チェックへの影響なし）。
    """
    year, month_num = 0, 0
    ym = re.search(r'(\d{4})年(\d+)月', excel_text)
    if ym:
        year, month_num = int(ym.group(1)), int(ym.group(2))

    entries, current = {}, None
    for line in excel_text.split('\n'):
        dm = re.match(r'【(\d+/\d+\([月火水木金土日]\))】', line)
        if dm:
            current = dm.group(1)
            entries[current] = {'lunch': '', 'snack': '', 'ingredients': '',
                                 'snack_am': '', 'snack_pm': ''}
        elif current:
            if line.startswith('昼食:'):
                entries[current]['lunch'] = line[3:].strip()
            elif line.startswith('午前おやつ:'):
                entries[current]['snack_am'] = line[6:].strip()
            elif line.startswith('午後おやつ:'):
                entries[current]['snack_pm'] = line[6:].strip()
            elif line.startswith('おやつ:'):
                entries[current]['snack'] = line[4:].strip()
            elif line.startswith('材料:'):
                entries[current]['ingredients'] = line[3:].strip()

    sorted_dates = sorted(
        entries.keys(),
        key=lambda ds: (_parse_date(ds, year) or datetime.date.max)
    )
    return year, month_num, sorted_dates, entries


# ─────────────────────────────────────────────
# チェック用定数
# ─────────────────────────────────────────────

# 2日連続チェック免除（完全一致）
_EXEMPT_EXACT = {
    '白米', '七分付き米', 'しょうが', '生姜', 'にんにく', 'みそ', '味噌', '酢', '白ごま',
    '人参', '玉ねぎ', '鶏肉', '豚肉', '牛肉', 'ひき肉', '牛乳',
    '昆布', 'かつお',  # 天然だし(かつお・昆布) をカッコ分割した際の破片を免除
    '昆布出し',  # 毎日使用OK（漢字表記のため'だし'部分一致に引っかからないため明示）
}
# 2日連続チェック免除（部分一致：これを含むトークンは免除）
# ※ 材料欄の表記は園によってひらがな／漢字が混在する（ごま/胡麻、しょうが/生姜等）ため、
#   両方の表記を登録しておく
_EXEMPT_SUB = ['醤油', '砂糖', 'みりん', '酒', '塩', '油',
               'だし', '出し', 'ごま', '胡麻', 'しょうが', '生姜', '水', '片栗粉', '小麦粉']
# 調味料独自チェックで管理するもの（汎用2日連続から除外）
_SEASONING_HANDLED = {'シャンタン', '中華だし', 'コンソメ', 'カレー', 'ソース', 'チーズ'}


def _is_exempt(token):
    if token in _EXEMPT_EXACT or token in _SEASONING_HANDLED:
        return True
    return any(s in token for s in _EXEMPT_SUB)


SEASONING_2DAY = ['シャンタン', '中華だし', 'コンソメ', 'カレー', 'ソース']
MEAT_3DAY      = ['豚肉', '鶏肉', '牛肉', 'ひき肉']
IMO_KW         = ['じゃが芋', 'さつま芋', '里芋', 'かぼちゃ']
NERIMONO_KW    = ['ちくわ', 'かにかま', '赤かまぼこ']
FISH_KW        = ['サケ', 'サーモン', 'サバ', 'サワラ', 'タラ', 'アジ', 'イワシ',
                  'ブリ', 'カレイ', 'メカジキ', 'タイ', 'マグロ', 'カツオ',
                  'シシャモ', 'ほっけ', '白身魚']
FISH_WITH_TUNA = FISH_KW + ['ツナ']
TOFU_KW        = ['木綿豆腐', '焼き豆腐', '厚揚げ', '油揚げ', '高野豆腐']
NOODLE_KW      = ['スパゲティ', 'うどん', 'めん', '麺', 'そば', '丼', 'ラーメン']
MUSHROOM_KW    = ['しめじ', 'エリンギ', 'えのき', 'なめこ']
MON_NG_ITEMS   = ['ロールパン', '食パン', 'りんご', 'バナナ', 'オレンジ',
                  '太もやし', '切干大根', 'ひじき', '高野豆腐']
PREP_KW        = [
    '玉ねぎ', '人参', '白菜', '焼き豆腐', 'かぼちゃ', '大根', 'いんげん',
    'えのき', 'キャベツ', 'じゃが芋', '里芋', '厚揚げ', 'しめじ', 'れんこん',
    '木綿豆腐', 'なめこ', 'エリンギ', 'ちくわ', '板こんにゃく', '糸こんにゃく',
    'さつま芋', 'ごぼう', 'ささがきごぼう', 'にら', 'マッシュルーム',
    'きゅうり', 'トマト', 'なす', '春雨', '切干大根', 'ロースハム',
    'ウインナー', 'たけのこ', 'アスパラガス',
]
FRUIT_KW       = [
    '果物', 'みかん', 'りんご', 'バナナ', 'オレンジ', 'パイン', 'パイナップル',
    'もも', '桃', 'ぶどう', 'ブドウ', 'いちご', 'イチゴ', 'メロン', 'すいか',
    'スイカ', 'キウイ', 'なし', '梨', 'マンゴー', 'さくらんぼ',
]
# 果物の種類・有無が日によって変わり得る総称デザート名。献立名にこれらの語が
# あるだけでは、具体的な果物名を書かない運用が正当（例：「手作りゼリー」の
# 中身がオレンジ果汁でも、献立名は「ゼリー」のまま）。「材料に果物があるのに
# 献立名に記載がない」チェック（reverse_naming_check）の母集団から除外する。
FRUIT_OPTIONAL_DESSERT_KW = ['ゼリー', 'プリン', 'ムース', '寒天', 'ババロア', 'シャーベット']
# 献立名の特定食材と材料欄の食い違い（表記ゆれ・別の食材へのすり替わり）を
# 決定的に判定するための同義語グループ。ここに載っている語は「献立名にあれば
# 材料欄にも（同義語のいずれかで）存在するはず」という前提で照合する。
# 以前はAIの自由記述チェックに任せていたが、言い回し次第で幻覚（実在するのに
# 「ない」と誤診断）を起こしていたため廃止し、この決定的グループ照合に一本化した。
# 新しい食材で同種の指摘が出た場合は、ここにグループを追加するだけで対応できる。
FOOD_SYNONYM_GROUPS = [
    ['パイン', 'パイナップル'],
    ['もも', '桃'],
    ['いちご', 'イチゴ'],
    ['ぶどう', 'ブドウ'],
    ['すいか', 'スイカ'],
    ['なし', '梨'],
    ['みかん'],
    ['りんご'],
    ['バナナ'],
    ['オレンジ', 'マーマレード'],  # 「オレンジ蒸しパン」等はマーマレードがあればOK
    ['メロン'],
    ['キウイ'],
    ['マンゴー'],
    ['さくらんぼ'],
    ['チンゲン菜', '青梗菜'],
    ['牛乳', 'ぎゅうにゅう'],
]

# 果物等の名前を含むが、実際にはその風味と無関係な慣用的な商品名・料理名。
# 献立名と材料の照合チェックで誤って「材料に見当たらない」と指摘しないよう、
# 判定前にこれらの語を献立名テキストから取り除く。
FOOD_NAME_IDIOMS = ['メロンパン']


def reverse_naming_check(kw_group, label, sorted_dates, lunch, snack, ing, day_ngs,
                          min_others=2, ratio=0.85):
    """「材料欄にはあるが献立名に記載がない」表記漏れを、ファイル内の運用パターンから判定する。
    対象日own自身を母集団から除いた残りの日（leave-one-out）で「普段は献立名に明記する」
    割合を求め、それが高い園に限って外れ値（記載漏れ）を報告する。
    普段から献立名に書かない園（割合が低い園）では何も報告しない。
    サンプルが少数（min_others未満）の場合は判定を保留する（自己参照による希釈を避けるため）。"""
    material_days = [ds for ds in sorted_dates if any(k in ing(ds) for k in kw_group)]
    named_days = {ds for ds in material_days if any(k in (lunch(ds) + ' ' + snack(ds)) for k in kw_group)}
    for ds in material_days:
        if ds in named_days:
            continue
        others = [d for d in material_days if d != ds]
        if len(others) < min_others:
            continue
        if sum(d in named_days for d in others) / len(others) >= ratio:
            day_ngs[ds].append(f'● 材料に「{label}」があるが献立名に記載がない（表記漏れの可能性）')


def bare_item_forward_check(bare_kw, label, sorted_dates, lunch, snack, ing, day_ngs,
                             min_others=2, ratio=0.85):
    """「献立名に○○が単体の項目として出ているのに材料欄にない」を、
    ファイル内の運用パターンから判定する（reverse_naming_checkと対称の考え方）。
    「フルーツヨーグルト」のような複合料理名は対象外（常時必須のチェックは別途行う）。
    対象日own自身を母集団から除いた残りの日（leave-one-out）で「普段は材料欄に明記する」
    割合を求め、それが高い園に限って記載漏れを報告する。普段から単体提供時は材料欄に
    書かない運用の園（割合が低い園）では何も報告しない。"""
    def _bare_items(text):
        return [re.sub(r'[（(].*?[）)]', '', i).strip() for i in text.split('/')]

    bare_days = [
        ds for ds in sorted_dates
        if bare_kw in _bare_items(lunch(ds) + '/' + snack(ds))
    ]
    has_material_days = {ds for ds in bare_days if bare_kw in ing(ds)}
    for ds in bare_days:
        if ds in has_material_days:
            continue
        others = [d for d in bare_days if d != ds]
        if len(others) < min_others:
            continue
        if sum(d in has_material_days for d in others) / len(others) >= ratio:
            day_ngs[ds].append(f'● 「{label}」単体提供があるが材料に{label}なし（記載漏れの可能性）')


def get_github_token():
    try:
        return st.secrets.get("GITHUB_TOKEN", "")
    except Exception:
        return ""


def push_rules_to_github(text):
    token = get_github_token()
    if not token:
        return False, "GitHubトークンが未設定です"

    api_url = (
        f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
        f"/contents/{GITHUB_RULES_PATH}"
    )
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json",
    }

    # 現在のSHAを取得
    req = urllib.request.Request(
        api_url + f"?ref={GITHUB_BRANCH}", headers=headers
    )
    try:
        with urllib.request.urlopen(req) as resp:
            sha = json.loads(resp.read())["sha"]
    except Exception as e:
        return False, f"GitHub取得エラー: {e}"

    # ファイルを更新
    payload = json.dumps({
        "message": "ルール管理から更新",
        "content": base64.b64encode(text.encode("utf-8")).decode("utf-8"),
        "sha": sha,
        "branch": GITHUB_BRANCH,
    }).encode("utf-8")

    req = urllib.request.Request(api_url, data=payload, headers=headers, method="PUT")
    try:
        with urllib.request.urlopen(req):
            return True, "保存しました（GitHub反映済み）"
    except Exception as e:
        return False, f"GitHub更新エラー: {e}"


def push_rules_list_to_github(rules_list):
    """複数ルール（rules.json）をGitHubに保存。"""
    token = get_github_token()
    if not token:
        return False, "GitHubトークンが未設定です"

    api_url = (
        f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO}"
        f"/contents/{GITHUB_RULES_JSON_PATH}"
    )
    headers = {
        "Authorization": f"token {token}",
        "Accept": "application/vnd.github.v3+json",
        "Content-Type": "application/json",
    }

    sha = None
    req = urllib.request.Request(
        api_url + f"?ref={GITHUB_BRANCH}", headers=headers
    )
    try:
        with urllib.request.urlopen(req) as resp:
            sha = json.loads(resp.read())["sha"]
    except Exception:
        pass  # 新規ファイルの場合はSHAなし

    content_text = json.dumps({"rules": rules_list}, ensure_ascii=False, indent=2)
    body = {"message": "ルール管理から更新", "branch": GITHUB_BRANCH,
            "content": base64.b64encode(content_text.encode("utf-8")).decode("utf-8")}
    if sha:
        body["sha"] = sha

    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(api_url, data=payload, headers=headers, method="PUT")
    try:
        with urllib.request.urlopen(req):
            return True, "保存しました（GitHub反映済み）"
    except Exception as e:
        return False, f"GitHub更新エラー: {e}"


# ─────────────────────────────────────────────
# 多形式Excelパーサー（さかえ保育園・おおみや・ゆめのはな対応）
# ─────────────────────────────────────────────

def _detect_sheet_format(df):
    """シートのフォーマット種別を返す: 'sakae' / 'omiya' / 'mebaenomori' / 'yumehana' / 'yamazaki' / 'kunimi' / 'default'"""
    all_text = ' '.join(str(v) for v in df.values.flatten() if pd.notna(v))
    if '熱と力になるもの' in all_text:
        return 'sakae'
    # 北野田こども園様形式（サイクル献立）：'材料名'+'献立名'（omiya）や'離乳食メニュー'
    # （tomikiya）の一般判定より先に、学校名で確実に判定する
    if '北野田こども園' in all_text and '離乳食メニュー' in all_text:
        return 'kitanoda_baby'
    if '北野田こども園' in all_text:
        return 'kitanoda'
    if '完了期' in all_text and '後期' in all_text:  # 岩戸こども園様形式（離乳食・後期/完了期の2段おやつ・園名の記載なし）
        return 'iwato'
    if '◎は10時おやつ' in all_text or ('材料名' in all_text and '献立名' in all_text):
        return 'omiya'
    if 'つかみ食べ練習用野菜' in all_text:  # 美山保育園形式（月別シート・datetime日付・おやつcol10）
        return 'miyama'
    # 「うどんの日以外はおかゆ」は富喜屋テンプレート系で共通に使われる文言のため、
    # 「離乳食メニュー」（tomikiya形式の目印）がある場合はそちらを優先する
    if 'うどんの日以外はおかゆ' in all_text and '離乳食メニュー' not in all_text:  # めばえの森（yumehanaより先に判定）
        return 'mebaenomori'
    if 'さかえ保育園' in all_text and '離乳食' in all_text:  # さかえ保育園 離乳食形式（2週間ローテーション・「日・日」表記）
        return 'sakae_baby'
    if '京町堀バンビ園' in all_text:  # 京町堀バンビ園形式（離乳食・Excelシリアル日付・おやつ単一列）
        return 'kyomachibori'
    if '離乳食メニュー' in all_text:  # 富喜屋提供の離乳食テンプレート（ぴよ・ぴよ保育園、鴻池第二バンビ等・複数園で共通）
        return 'tomikiya'
    if '初期には入りません' in all_text or 'おかゆが付きます' in all_text:
        return 'yumehana'
    if '初期・アレルギーには' in all_text:  # 歩学園バンビ形式（離乳食・datetime日付・おやつcol8-9）
        return 'ayumi'
    if 'くにみ子ども園' in all_text:  # くにみ子ども園形式（横並び・4列/日・午前/昼食/午後・乳児幼児2分量）
        return 'kunimi'
    # 山崎幼稚園形式（横並び・3列/日）：セル単体が「材料表」と完全一致する場合のみ判定。
    # 部分一致にすると「献立材料表」のようなタイトル文言を含む他園（例：東久留米おひさま）を
    # 誤って山崎形式と判定してしまう（列オフセットが異なるため誤爆する）。
    cell_values = {str(v).strip() for v in df.values.flatten() if pd.notna(v)}
    if '材料表' in cell_values:
        return 'yamazaki'
    return 'default'


def _extract_year_month(df, fname=""):
    """先頭10行から年月を抽出。タイトルが右端の列にあるフォーマットもあるため
    列は全て走査する（行のみ先頭10行に限定）。(year_int, month_int, label_str)
    シート内に年月表記がない場合（月が複数シートに分かれ、2シート目以降には
    タイトルを繰り返さない形式等）は、アップロードファイル名からのフォールバック
    抽出を試みる（例：「2026年9月　◯◯こども園様.xls」）。全角数字（「９月」等）
    にも対応するため、判定前にNFKC正規化で半角化する。年の記載が全くなく
    月のみの場合（例：「９月普通献立.xls」）は、実行時点の日付から年を推定する
    （対象月が現在月以降ならその年、過去月なら翌年とみなす）。"""
    n_rows, n_cols = df.shape
    for r in range(min(10, n_rows)):
        for c in range(n_cols):
            v = unicodedata.normalize('NFKC', str(df.iloc[r, c]).strip())
            m = re.match(r'(\d{4})年(\d{1,2})月', v)
            if m:
                y, mo = int(m.group(1)), int(m.group(2))
                return y, mo, f"{y}年{mo:02d}月"
    if fname:
        fname_n = unicodedata.normalize('NFKC', fname)
        m = re.search(r'(\d{4})年(\d{1,2})月', fname_n)
        if m:
            y, mo = int(m.group(1)), int(m.group(2))
            return y, mo, f"{y}年{mo:02d}月"
        m = re.search(r'(\d{1,2})月', fname_n)
        if m:
            mo = int(m.group(1))
            if 1 <= mo <= 12:
                today = datetime.date.today()
                y = today.year if mo >= today.month else today.year + 1
                return y, mo, f"{y}年{mo:02d}月"
    return 0, 0, ""


def _excel_to_text_sakae(df, fname=""):
    """さかえ保育園形式（縦並び・4列材料）→ 構造化テキスト"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    year_num, month_num, year_month = _extract_year_month(df, fname)

    # 「昼食」から始まる保育園形式に加え、児童養護施設等の「朝食」から始まる
    # 3食（朝食/昼食/夕食）形式にも対応する。おやつ以外の食事区分はすべて
    # 「昼食」バケットにまとめて格納する（チェック側は昼食/おやつの2分類のため）。
    _MEAL_SECTION_LABELS = {'朝食', '昼食', '夕食'}

    days = []
    cur = None

    for r in range(n_rows):
        col0, col1, col2 = cv(r, 0), cv(r, 1), cv(r, 2)

        if re.match(r'^\d{1,2}(\.0)?$', col0) and col1 in _MEAL_SECTION_LABELS:
            if cur is not None:
                days.append(cur)
            cur = {'day': int(float(col0)), 'dow': '?', 'lunch': [], 'snack': [], 'mats': [], 'in_snack': False}
            if col2:
                cur['lunch'].append(col2)
            for c in range(3, min(7, n_cols)):  # 列3-6が材料、列7はエネルギー値
                v = cv(r, c)
                if v:
                    cur['mats'].append(v)

        elif re.match(r'^[月火水木金土日]$', col0) and cur is not None:
            cur['dow'] = col0
            # 曜日行に献立名・材料が同居する場合（「土」行に鶏肉となすのみそ炒め等）
            if col2 and col1 == '':
                if cur['in_snack']:
                    cur['snack'].append(col2)
                else:
                    cur['lunch'].append(col2)
                for c in range(3, min(7, n_cols)):
                    v = cv(r, c)
                    if v:
                        cur['mats'].append(v)

        elif col1 == '午後おやつ' and cur is not None:
            cur['in_snack'] = True
            if col2:
                cur['snack'].append(col2)
            for c in range(3, min(7, n_cols)):
                v = cv(r, c)
                if v:
                    cur['mats'].append(v)

        elif col1 in _MEAL_SECTION_LABELS and cur is not None:
            # 同じ日の別の食事区分への切り替え（朝食→昼食→夕食）。おやつではない。
            cur['in_snack'] = False
            if col2:
                cur['lunch'].append(col2)
            for c in range(3, min(7, n_cols)):
                v = cv(r, c)
                if v:
                    cur['mats'].append(v)

        elif cur is not None and col2 and col0 == '' and col1 == '':
            if cur['in_snack']:
                cur['snack'].append(col2)
            else:
                cur['lunch'].append(col2)
            for c in range(3, min(7, n_cols)):
                v = cv(r, c)
                if v:
                    cur['mats'].append(v)

    if cur is not None:
        days.append(cur)

    lines = []
    if year_month:
        lines += [f"# 献立データ {year_month}", ""]

    for d in days:
        label = f"{month_num}/{d['day']}({d['dow']})" if month_num else f"?/{d['day']}({d['dow']})"
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['snack']:
            lines.append(f"おやつ: {' / '.join(d['snack'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_omiya(df, fname=""):
    """おおみやこども園形式（縦並び・1セル全材料）→ 構造化テキスト"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    year_num, month_num, year_month = _extract_year_month(df, fname)

    def clean_mat(text):
        text = re.sub(r'\(\d+g\)', '', text)         # 量(30g)を除去
        text = text.replace('*', '').replace('＊', '') # *印を除去
        text = re.sub(r'[／\n]', ',', text)           # ／と改行をカンマに
        return text

    def clean_oyatsu(text):
        text = text.replace('◎', '').replace('＊', '')
        text = re.sub(r'\(\d+g\)', '', text)
        return text.replace('\n', ' / ').strip()

    days = []

    for r in range(n_rows):
        col0 = cv(r, 0)
        col1 = cv(r, 1)
        col2 = cv(r, 2)

        if re.match(r'^\d{1,2}(\.0)?$', col0):
            day_int = int(float(col0))
            lunch_names = [x.strip() for x in re.split(r'[,\n]', col1) if x.strip()] if col1 else []
            mats_raw = clean_mat(col2) if col2 else ""
            mat_list = [x.strip() for x in re.split(r'[,、]', mats_raw) if x.strip()]

            oyatsu_names = []
            col5 = cv(r, 5) if n_cols > 5 else ""
            if col5:
                oyatsu_text = clean_oyatsu(col5)
                oyatsu_names = [x.strip() for x in re.split(r'[\s/／,、\n]+', oyatsu_text) if x.strip()]

            days.append({'day': day_int, 'dow': '?', 'lunch': lunch_names, 'snack': oyatsu_names, 'mats': mat_list})

        elif re.match(r'^[月火水木金土日]$', col0) and days:
            days[-1]['dow'] = col0

    lines = []
    if year_month:
        lines += [f"# 献立データ {year_month}", ""]

    for d in days:
        label = f"{month_num}/{d['day']}({d['dow']})" if month_num else f"?/{d['day']}({d['dow']})"
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['snack']:
            lines.append(f"おやつ: {' / '.join(d['snack'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_yumehana(df, fname=""):
    """ゆめのはなこども園形式（週別シート・Excelシリアル日付）→ 構造化テキスト"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    def clean_mat(v):
        return v.replace('＊', '').replace('*', '').strip()

    year_num, month_num, year_month = _extract_year_month(df, fname)

    days = {}   # label → {'lunch', 'snack', 'mats'}
    day_order = []

    for r in range(n_rows):
        col0, col1, col2 = cv(r, 0), cv(r, 1), cv(r, 2)

        if re.match(r'^\d{5}(\.0)?$', col0):  # Excelシリアル番号（5桁）
            serial = int(float(col0))
            dow = col1 if re.match(r'^[月火水木金土日]$', col1) else '?'
            try:
                d = datetime.date(1899, 12, 30) + datetime.timedelta(days=serial)
                label = f"{d.month}/{d.day}({dow})"
                if not year_num:
                    year_num, month_num = d.year, d.month
            except Exception:
                label = f"?/?({dow})"

            if label not in days:
                days[label] = {'lunch': [], 'snack': [], 'mats': []}
                day_order.append(label)

            if col2 and '印は初期' not in col2 and 'おかゆが付きます' not in col2:
                days[label]['lunch'].append(col2)
            for c in range(3, n_cols):
                v = cv(r, c)
                if v:
                    days[label]['mats'].append(clean_mat(v))

        elif col0 == '' and col2 and day_order:
            if '印は初期' not in col2 and 'おかゆが付きます' not in col2:
                label = day_order[-1]
                days[label]['lunch'].append(col2)
                for c in range(3, n_cols):
                    v = cv(r, c)
                    if v:
                        days[label]['mats'].append(clean_mat(v))

    lines = []
    if year_num and month_num:
        lines += [f"# 献立データ {year_num}年{month_num:02d}月", ""]

    for label in day_order:
        d = days[label]
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['snack']:
            lines.append(f"おやつ: {' / '.join(d['snack'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_mebaenomori(df):
    """めばえの森保育園形式（週別シート・datetime文字列日付・離乳食・2行/日）→ 構造化テキスト"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    def clean_mat(v):
        return v.replace('＊', '').replace('*', '').strip()

    _NOTE_PHRASES = ('初期には入りません', 'おかゆが付きます', '富喜屋', '株式会社')
    _HOLIDAY_KW   = ('山の日', '海の日', '振替休日', '祝日', '休園', '夏期休暇',
                     '冬期休暇', '春期休暇', 'こどもの日', '天皇誕生日')

    year_num, month_num = 0, 0

    days = {}     # label → {'lunch': [], 'mats': []}
    day_order = []

    for r in range(n_rows):
        col0 = cv(r, 0)
        col1 = cv(r, 1)
        col2 = cv(r, 2)

        # 日付行: pandasがdtype=strで読むと 'YYYY-MM-DD HH:MM:SS' 形式になる
        dm = re.match(r'^(\d{4})-(\d{2})-(\d{2})', col0)
        if dm:
            y, mo, day = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
            if not year_num:
                year_num, month_num = y, mo
            dow = col1 if re.match(r'^[月火水木金土日]$', col1) else '?'
            label = f"{mo}/{day}({dow})"

            if label not in days:
                days[label] = {'lunch': [], 'mats': []}
                day_order.append(label)

            # 祝日・休園は料理名として追加しない
            if col2 and not any(h in col2 for h in _HOLIDAY_KW):
                days[label]['lunch'].append(col2)
            for c in range(3, n_cols):
                v = clean_mat(cv(r, c))
                if v:
                    days[label]['mats'].append(v)

        # 2行目（col0・col1ともに空）: 同日の材料継続行または副料理行
        elif col0 == '' and col1 == '' and day_order:
            # 行全体のテキストに注記フレーズがあればスキップ
            row_text = ' '.join(cv(r, c) for c in range(n_cols))
            if any(p in row_text for p in _NOTE_PHRASES):
                continue
            mats_in_row = [clean_mat(cv(r, c)) for c in range(3, n_cols)
                           if clean_mat(cv(r, c))]
            if col2 or mats_in_row:
                label = day_order[-1]
                if col2:
                    days[label]['lunch'].append(col2)
                days[label]['mats'].extend(mats_in_row)

    lines = []
    if year_num and month_num:
        lines += [f"# 献立データ {year_num}年{month_num:02d}月", ""]

    for label in day_order:
        d = days[label]
        if not d['lunch'] and not d['mats']:
            continue  # 祝日・空日はスキップ
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_yamazaki(df, fname=""):
    """山崎幼稚園形式（横並び・2ブロック・3列/日・日付col_c/献立+材料col_c+1）→ 構造化テキスト"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    year_num, month_num, year_month = _extract_year_month(df, fname)

    _SKIP = {"[昼]", "[午後]", "献立名", "材料", "材料表", "日付", "区分",
             "お箸もいります", "おはしもいります"}

    def is_valid(v):
        if not v or v in _SKIP:
            return False
        if v.startswith("※") or v.startswith("【"):
            return False
        try:
            float(v)   # 数量（g数等）は除外
            return False
        except ValueError:
            pass
        return True

    # 日付行を検索（「N日(曜)」パターンが3個以上ある行）
    blocks = []
    for r in range(n_rows):
        temp = {}
        for c in range(n_cols):
            v = cv(r, c)
            if re.match(r'^\d+日\([月火水木金土日]\)$', v):
                temp[c] = v
        if len(temp) >= 3:
            blocks.append((r, temp))

    lines = []
    if year_month:
        lines += [f"# 献立データ {year_month}", ""]

    for block_idx, (date_row, date_cols) in enumerate(blocks):
        block_end = blocks[block_idx + 1][0] if block_idx + 1 < len(blocks) else n_rows

        # 材料表開始行（col 1 が "材料表" または "材料"）
        mat_row = None
        for r in range(date_row + 1, block_end):
            if cv(r, 1) in ("材料表", "材料"):
                mat_row = r
                break

        dish_end = mat_row if mat_row is not None else block_end

        for col_c in sorted(date_cols.keys()):
            raw_date = date_cols[col_c]
            dm = re.match(r'(\d+)日\(([月火水木金土日])\)', raw_date)
            if dm and month_num:
                date_label = f"{month_num}/{dm.group(1)}({dm.group(2)})"
            else:
                date_label = raw_date

            # [午後]マーカーを col_c で検索（行は日によって異なる）
            afternoon_start = dish_end
            for r in range(date_row + 1, dish_end):
                if cv(r, col_c) == "[午後]":
                    afternoon_start = r
                    break

            # 昼食・おやつは col_c+1 から読む
            lunch, snack = [], []
            for r in range(date_row + 1, afternoon_start):
                v = cv(r, col_c + 1)
                if is_valid(v):
                    lunch.append(v)
            for r in range(afternoon_start, dish_end):
                v = cv(r, col_c + 1)
                if is_valid(v):
                    snack.append(v)

            # 材料は col_c+1 から読む（材料表セクション）
            mats = []
            if mat_row is not None:
                for r in range(mat_row, block_end):
                    v = cv(r, col_c + 1)
                    if is_valid(v):
                        mats.append(v)

            lines.append(f"【{date_label}】")
            if lunch:
                lines.append(f"昼食: {' / '.join(lunch)}")
            if snack:
                lines.append(f"おやつ: {' / '.join(snack)}")
            if mats:
                lines.append(f"材料: {', '.join(mats)}")
            lines.append("")

    return '\n'.join(lines)


def _excel_to_text_ayumi(df):
    """歩学園バンビ形式（週別シート・datetime日付・離乳食・おやつcol8-9）→ 構造化テキスト"""
    n_rows, n_cols = df.shape
    _DOW = '月火水木金土日'

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    def clean_mat(v):
        return v.replace('＊', '').replace('*', '').strip()

    _NOTE = ('初期・アレルギーには', '麺の日以外', '初期におやつ', 'お菓子は、',
             '完了期のアレルギー', '富喜屋', '株式会社')

    year_num, month_num = 0, 0
    days = {}
    day_order = []

    def _parse_snack(text):
        """'料理名\n（材料1、材料2）' → (name, [mats])"""
        if not text:
            return None, []
        text = re.sub(r'\s+', ' ', text.replace('\n', ' '))
        m = re.search(r'[（(]([^）)]+)[）)]', text)
        if m:
            name = text[:m.start()].strip()
            ings = [clean_mat(s.strip()) for s in re.split(r'[,、\s]+', m.group(1)) if s.strip()]
        else:
            name = text.strip()
            ings = []
        return name or None, ings

    for r in range(n_rows):
        col0 = cv(r, 0)
        col2 = cv(r, 2)

        # 日付行（pandasがdatetime→'YYYY-MM-DD HH:MM:SS'に変換）
        dm = re.match(r'^(\d{4})-(\d{2})-(\d{2})', col0)
        if dm:
            y, mo, day_i = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
            if not year_num:
                year_num, month_num = y, mo
            # 曜日を日付から計算（col1が空のケースを補完）
            try:
                dow = _DOW[datetime.date(y, mo, day_i).weekday()]
            except Exception:
                dow = '?'
            label = f"{mo}/{day_i}({dow})"
            if label not in days:
                days[label] = {'lunch': [], 'snack': [], 'mats': []}
                day_order.append(label)

            if col2 and not any(n in col2 for n in _NOTE):
                days[label]['lunch'].append(col2)
            for c in range(3, min(8, n_cols)):
                v = clean_mat(cv(r, c))
                if v and not any(n in v for n in _NOTE):
                    days[label]['mats'].append(v)

            # おやつ（col8: 中期～後期、col9: 完了期）
            for sc in (8, 9):
                sv = cv(r, sc)
                if not sv:
                    continue
                sname, sings = _parse_snack(sv)
                if sname and sname not in days[label]['snack']:
                    days[label]['snack'].append(sname)
                days[label]['mats'].extend(sings)

        # 継続行（col0が空）
        elif col0 == '' and day_order:
            row_text = ' '.join(cv(r, c) for c in range(n_cols))
            if any(n in row_text for n in _NOTE):
                continue
            label = day_order[-1]
            if col2:
                days[label]['lunch'].append(col2)
            for c in range(3, min(8, n_cols)):
                v = clean_mat(cv(r, c))
                if v and not any(n in v for n in _NOTE):
                    days[label]['mats'].append(v)

    lines = []
    if year_num and month_num:
        lines += [f"# 献立データ {year_num}年{month_num:02d}月", ""]

    for label in day_order:
        d = days[label]
        if not d['lunch'] and not d['mats']:
            continue
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['snack']:
            lines.append(f"おやつ: {' / '.join(d['snack'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_tomikiya(df):
    """富喜屋提供の離乳食テンプレート形式（ぴよ・ぴよ保育園、鴻池第二バンビ等・複数園共通）。
    1日が複数行（主菜＋野菜スープ［＋果物］）に分かれ、日付はブロック先頭行のみ・
    後続行は日付欄が空のまま材料が続く。ヘッダー行の「献立名」「おやつ」セル位置を
    動的検出して材料欄の範囲を決める（列数が多少変わっても追従できる設計）。"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    header_row = name_col = None
    for r in range(n_rows):
        for c in range(n_cols):
            if cv(r, c) == '献立名':
                header_row, name_col = r, c
                break
        if header_row is not None:
            break
    if header_row is None:
        return ""

    snack_col = n_cols
    for c in range(n_cols):
        if cv(header_row, c) == 'おやつ':
            snack_col = c
            break

    date_col = 0
    for c in range(n_cols):
        if '日付' in cv(header_row, c):
            date_col = c
            break

    mat_start, mat_end = name_col + 1, snack_col
    _DOW = ['月', '火', '水', '木', '金', '土', '日']
    _NOTE = ('初期', '完了期', 'お菓子は', '株式会社', '主食は', 'アレルギー',
             '離乳食は', '大きさ等', '普通食を刻んで')

    day_entries = []
    cur = None
    for r in range(header_row + 1, n_rows):
        date_raw = cv(r, date_col)
        col1_raw = cv(r, 1) if n_cols > 1 else ""
        if date_raw.startswith(('※', '*印')) or col1_raw.startswith(('※', '*印')):
            break  # 注記・免責文・提供元表記の行に到達したら終了

        row_text = ' '.join(cv(r, c) for c in range(n_cols))
        if any(n in row_text for n in _NOTE):
            continue  # 注記・免責文・提供元表記の行（日付欄の外側に出現するケース）は無視

        dm = re.match(r'(\d{4})-(\d{2})-(\d{2})', date_raw)
        serial_m = None if dm else re.match(r'^(\d{5})(\.0)?$', date_raw)
        name_v = cv(r, name_col)
        mats_v = [cv(r, c) for c in range(mat_start, mat_end) if cv(r, c)]
        snack_v = cv(r, snack_col) if snack_col < n_cols else ""

        if dm:
            y, mo, d = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
            dow = _DOW[datetime.date(y, mo, d).weekday()]
            cur = {'year': y, 'month': mo, 'day': d, 'dow': dow,
                   'lunch': [], 'mats': [], 'snack': ''}
            day_entries.append(cur)
        elif serial_m:
            # 日付欄がExcelシリアル番号（未整形の数値）で入っている書式に対応
            dt = datetime.date(1899, 12, 30) + datetime.timedelta(days=int(serial_m.group(1)))
            y, mo, d, dow = dt.year, dt.month, dt.day, _DOW[dt.weekday()]
            cur = {'year': y, 'month': mo, 'day': d, 'dow': dow,
                   'lunch': [], 'mats': [], 'snack': ''}
            day_entries.append(cur)
        elif cur is None or not (name_v or mats_v or snack_v):
            continue  # 日付が確定する前、または内容のない行（週またぎの休園note等）は無視

        if name_v:
            cur['lunch'].append(name_v)
        cur['mats'].extend(mats_v)
        if snack_v:
            if not cur['snack']:
                cur['snack'] = snack_v
            else:
                # おやつ名の次行に、そのおやつの原材料が同じ列に入る形式（富喜屋テンプレート等）
                cur['mats'].append(snack_v)

    if not day_entries:
        return ""

    y0, mo0 = day_entries[0]['year'], day_entries[0]['month']
    lines = [f"# 献立データ {y0}年{mo0:02d}月", ""]
    for e in day_entries:
        lines.append(f"【{e['month']}/{e['day']}({e['dow']})】")
        if e['lunch']:
            lines.append(f"昼食: {' / '.join(e['lunch'])}")
        if e['snack']:
            lines.append(f"おやつ: {e['snack']}")
        if e['mats']:
            lines.append(f"材料: {', '.join(e['mats'])}")
        lines.append("")
    return "\n".join(lines)


def _excel_to_text_kyomachibori(df):
    """京町堀バンビ園形式（離乳食・週別シート・Excelシリアル日付・おやつ単一列）→ 構造化テキスト"""
    n_rows, n_cols = df.shape
    _DOW = '月火水木金土日'

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    def clean_mat(v):
        return v.replace('＊', '').replace('*', '').strip()

    _NOTE = ('初期', '完了期', 'お菓子は', '株式会社', '主食は', 'アレルギー', 'gです',
             '離乳食は', '大きさ等', '普通食を刻んで')

    # ヘッダー行から「おやつ」列（材料欄の終端）を特定
    snack_col = None
    for r in range(min(5, n_rows)):
        for c in range(n_cols):
            if cv(r, c) == 'おやつ':
                snack_col = c
                break
        if snack_col is not None:
            break
    mat_end_col = snack_col if snack_col is not None else n_cols

    year_num, month_num = 0, 0
    days = {}
    day_order = []

    for r in range(n_rows):
        col0, col2 = cv(r, 0), cv(r, 2)

        y = mo = day_i = None
        dm = re.match(r'^(\d{4})-(\d{2})-(\d{2})', col0)
        if dm:
            y, mo, day_i = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
        elif re.match(r'^\d{5}(\.0)?$', col0):
            try:
                d = datetime.date(1899, 12, 30) + datetime.timedelta(days=int(float(col0)))
                y, mo, day_i = d.year, d.month, d.day
            except Exception:
                pass

        if y:
            if not year_num:
                year_num, month_num = y, mo
            try:
                dow = _DOW[datetime.date(y, mo, day_i).weekday()]
            except Exception:
                dow = '?'
            label = f"{mo}/{day_i}({dow})"
            if label not in days:
                days[label] = {'lunch': [], 'snack': [], 'mats': []}
                day_order.append(label)

            if col2 and not any(n in col2 for n in _NOTE):
                days[label]['lunch'].append(col2)
            for c in range(3, mat_end_col):
                v = clean_mat(cv(r, c))
                if v and not re.match(r'^\d+(\.\d+)?$', v) and not any(n in v for n in _NOTE):
                    days[label]['mats'].append(v)
            if snack_col is not None:
                sv = clean_mat(cv(r, snack_col))
                if sv and sv != 'おやつ' and sv not in days[label]['snack']:
                    days[label]['snack'].append(sv)

        elif col0 == '' and day_order:
            row_text = ' '.join(cv(r, c) for c in range(n_cols))
            if any(n in row_text for n in _NOTE):
                continue
            label = day_order[-1]
            if col2:
                days[label]['lunch'].append(col2)
            for c in range(3, mat_end_col):
                v = clean_mat(cv(r, c))
                if v and not re.match(r'^\d+(\.\d+)?$', v) and not any(n in v for n in _NOTE):
                    days[label]['mats'].append(v)

    lines = []
    if year_num and month_num:
        lines += [f"# 献立データ {year_num}年{month_num:02d}月", ""]

    for label in day_order:
        d = days[label]
        if not d['lunch'] and not d['mats']:
            continue
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['snack']:
            lines.append(f"おやつ: {' / '.join(d['snack'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_sakae_baby(df):
    """さかえ保育園 離乳食形式（2週間ローテーション・「日・日」表記の日付）→ 構造化テキスト"""
    n_rows, n_cols = df.shape
    _DOW = '月火水木金土日'

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    def clean_mat(v):
        return v.replace('＊', '').replace('*', '').strip()

    _NOTE = ('初期には入りません', 'おかゆが付きます', '麦茶が付きます', '株式会社')

    # 年月抽出：「2026年度　8月」のような表記に対応
    year_num, month_num = 0, 0
    for r in range(min(5, n_rows)):
        for c in range(n_cols):
            m = re.search(r'(\d{4})年度\s*(\d{1,2})月', cv(r, c))
            if m:
                year_num, month_num = int(m.group(1)), int(m.group(2))
                break
        if year_num:
            break

    days = {}
    day_order = []

    def _get_or_create(label):
        if label not in days:
            days[label] = {'lunch': [], 'mats': []}
            day_order.append(label)
        return days[label]

    cur_labels = []  # 現在のブロックが対応する日付ラベル一覧（「N・M」で複数になる）

    for r in range(n_rows):
        col0, col1, col2 = cv(r, 0), cv(r, 1), cv(r, 2)
        row_text = ' '.join(cv(r, c) for c in range(n_cols))
        if any(n in row_text for n in _NOTE):
            continue

        if col0 and re.match(r'^[月火水木金土日]$', col1):
            day_nums = [int(x) for x in re.split(r'[・,、]', col0) if x.strip().isdigit()]
            cur_labels = []
            for dn in day_nums:
                dow = col1
                try:
                    if month_num:
                        dow = _DOW[datetime.date(year_num, month_num, dn).weekday()]
                except Exception:
                    pass
                label = f"{month_num}/{dn}({dow})"
                cur_labels.append(label)
                _get_or_create(label)

            if col2:
                for label in cur_labels:
                    days[label]['lunch'].append(col2)
            for c in range(3, n_cols):
                v = clean_mat(cv(r, c))
                if v:
                    for label in cur_labels:
                        days[label]['mats'].append(v)

        elif col0 == '' and col1 == '' and cur_labels:
            if col2:
                for label in cur_labels:
                    days[label]['lunch'].append(col2)
            for c in range(3, n_cols):
                v = clean_mat(cv(r, c))
                if v:
                    for label in cur_labels:
                        days[label]['mats'].append(v)

    lines = []
    if year_num and month_num:
        lines += [f"# 献立データ {year_num}年{month_num:02d}月", ""]

    for label in sorted(day_order, key=lambda s: int(re.match(r'\d+/(\d+)', s).group(1))):
        d = days[label]
        if not d['lunch'] and not d['mats']:
            continue
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _excel_to_text_miyama(df):
    """美山保育園形式（月別シート・datetime日付・離乳食・おやつcol10）→ 構造化テキスト"""
    n_rows, n_cols = df.shape
    _DOW = '月火水木金土日'

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    def clean_mat(v):
        return v.replace('＊', '').replace('*', '').strip()

    _NOTE = ('*印は初期には入りません', '麺の日以外の主食', '初期におやつ',
             'お菓子は、', '富喜屋', '株式会社', 'ここに日、祝',
             'つかみ食べ練習用野菜', '給食提供が無い日')

    year_num, month_num = 0, 0
    days = {}
    day_order = []

    def _parse_snack(text):
        if not text:
            return None, []
        text = re.sub(r'\s+', ' ', text.replace('\n', ' '))
        m = re.search(r'[（(]([^）)]+)[）)]', text)
        if m:
            name = text[:m.start()].strip()
            ings = [clean_mat(s.strip()) for s in re.split(r'[,、\s・]+', m.group(1)) if s.strip()]
        else:
            name = text.strip()
            ings = []
        return name or None, ings

    for r in range(n_rows):
        col0 = cv(r, 0)
        col2 = cv(r, 2)

        dm = re.match(r'^(\d{4})-(\d{2})-(\d{2})', col0)
        if dm:
            y, mo, day_i = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
            if not year_num:
                year_num, month_num = y, mo
            if y != year_num or mo != month_num:
                continue  # 別月（祝日リスト等）はスキップ
            try:
                dow = _DOW[datetime.date(y, mo, day_i).weekday()]
            except Exception:
                dow = '?'
            label = f"{mo}/{day_i}({dow})"
            if label not in days:
                days[label] = {'lunch': [], 'snack': [], 'mats': []}
                day_order.append(label)

            if col2 and not any(n in col2 for n in _NOTE):
                days[label]['lunch'].append(col2)
            for c in range(3, min(10, n_cols)):
                v = clean_mat(cv(r, c))
                if v and not any(n in v for n in _NOTE):
                    days[label]['mats'].append(v)

            # おやつ（col10: 中期〜後期）
            sv = cv(r, 10)
            if sv and not any(n in sv for n in _NOTE):
                sname, sings = _parse_snack(sv)
                if sname and sname not in days[label]['snack']:
                    days[label]['snack'].append(sname)
                days[label]['mats'].extend(sings)

        elif col0 == '' and day_order:
            row_text = ' '.join(cv(r, c) for c in range(n_cols))
            if any(n in row_text for n in _NOTE):
                continue
            label = day_order[-1]
            if col2 and not any(n in col2 for n in _NOTE):
                days[label]['lunch'].append(col2)
            for c in range(3, min(10, n_cols)):
                v = clean_mat(cv(r, c))
                if v and not any(n in v for n in _NOTE):
                    days[label]['mats'].append(v)

    lines = []
    if year_num and month_num:
        lines += [f"# 献立データ {year_num}年{month_num:02d}月", ""]

    for label in day_order:
        d = days[label]
        if not d['lunch'] and not d['mats']:
            continue
        lines.append(f"【{label}】")
        if d['lunch']:
            lines.append(f"昼食: {' / '.join(d['lunch'])}")
        if d['snack']:
            lines.append(f"おやつ: {' / '.join(d['snack'])}")
        if d['mats']:
            lines.append(f"材料: {', '.join(d['mats'])}")
        lines.append("")

    return '\n'.join(lines)


def _clean_kitanoda_mats(raw):
    """'／'・'、'区切りが混在する材料セルをトークンに分割し ', ' 区切りに統一。"""
    parts = re.split(r'[、／]+', raw)
    return ', '.join(p.strip() for p in parts if p.strip())


def _excel_to_text_kitanoda(df, fname=""):
    """北野田こども園様形式（幼児・サイクル献立）→ 構造化テキスト。
    1つの献立ブロックが4行（エネルギー/たんぱく質/脂質/食物繊維の各行）で構成され、
    同じ献立を使う日付（例：01日・15日・29日）は各行の日付欄に列挙される。
    材料は1セルに「／」区切りでまとめて入る（献立名の品数＋おやつの品数ぶんのグループ）。"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    year_num, month_num, year_month = _extract_year_month(df, fname)

    header_row = None
    for r in range(n_rows):
        if cv(r, 1) == '献立名':
            header_row = r
            break
    if header_row is None:
        return ""

    snack_col = 5 if n_cols > 5 else None
    for c in range(n_cols):
        if 'おやつ' in cv(header_row, c):
            snack_col = c
            break

    _DOW_SET = set('月火水木金土日')
    _DOW_ORDER = ['月', '火', '水', '木', '金', '土', '日']
    _DATE_RE = re.compile(r'^(\d{1,2})(\.0)?$')

    blocks = []
    r = header_row + 1
    while r < n_rows:
        col0 = cv(r, 0)
        if col0.startswith(('※', '◎')):
            break
        dm = _DATE_RE.match(col0)
        name_v = cv(r, 1)
        if not (dm and name_v):
            r += 1
            continue

        dates = [int(dm.group(1))]
        dow = None
        for rr in range(r + 1, min(r + 4, n_rows)):
            c0 = cv(rr, 0)
            if c0 in _DOW_SET:
                dow = c0
            else:
                dm2 = _DATE_RE.match(c0)
                if dm2:
                    dates.append(int(dm2.group(1)))

        blocks.append({
            'dates': dates, 'dow': dow,
            'lunch': name_v,
            'snack': cv(r, snack_col) if snack_col is not None else "",
            'mats': cv(r, 2),
        })
        r += 4

    if not blocks:
        return ""

    lines = []
    if year_month:
        lines += [f"# 献立データ {year_month}", ""]

    mats_cache = {}
    for e in blocks:
        mats_disp = mats_cache.get(e['mats'])
        if mats_disp is None:
            mats_disp = _clean_kitanoda_mats(e['mats'])
            mats_cache[e['mats']] = mats_disp
        lunch_disp = e['lunch'].replace(',', ' / ')
        snack_disp = e['snack'].replace(',', ' / ') if e['snack'] else ""

        for d in e['dates']:
            dow = e['dow']
            if not dow and year_num and month_num:
                try:
                    dow = _DOW_ORDER[datetime.date(year_num, month_num, d).weekday()]
                except Exception:
                    dow = '?'
            dow = dow or '?'
            label = f"{month_num}/{d}({dow})" if month_num else f"?/{d}({dow})"
            lines.append(f"【{label}】")
            lines.append(f"昼食: {lunch_disp}")
            if snack_disp:
                lines.append(f"おやつ: {snack_disp}")
            if mats_disp:
                lines.append(f"材料: {mats_disp}")
            lines.append("")

    return '\n'.join(lines)


def _excel_to_text_kitanoda_baby(df, fname=""):
    """北野田こども園様形式（離乳食・サイクル献立）→ 構造化テキスト。
    1日2行構成（1行目：日付Excelシリアル値＋主菜＋材料3列＋おやつ、
    2行目：副菜＋材料3列）。おやつ欄はセル内改行で「おやつ名\\n（材料）」の
    形式になっているため、名前と材料を分離して扱う。"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    header_row = None
    for r in range(n_rows):
        if cv(r, 0) == '月日':
            header_row = r
            break
    if header_row is None:
        return ""

    _DOW = ['月', '火', '水', '木', '金', '土', '日']
    _SERIAL_RE = re.compile(r'^(\d{5})(\.0)?$')

    day_entries = []
    cur = None
    for r in range(header_row + 1, n_rows):
        col0 = cv(r, 0)
        if col0.startswith(('*印', '◎')):
            break

        # 日付欄は「YYYY-MM-DD...」の日付文字列と、書式が数値のままの
        # Excelシリアル値の両方があり得るため両方に対応する
        ymd_m = re.match(r'(\d{4})-(\d{2})-(\d{2})', col0)
        serial_m = None if ymd_m else _SERIAL_RE.match(col0)
        name_v = cv(r, 1)
        mats_v = [cv(r, c) for c in (2, 3, 4) if cv(r, c)]
        snack_cell = cv(r, 5) if n_cols > 5 else ""

        if ymd_m:
            y, mo, d = int(ymd_m.group(1)), int(ymd_m.group(2)), int(ymd_m.group(3))
            cur = {'year': y, 'month': mo, 'day': d,
                   'dow': _DOW[datetime.date(y, mo, d).weekday()],
                   'lunch': [], 'mats': [], 'snack': ''}
            day_entries.append(cur)
        elif serial_m:
            dt = datetime.date(1899, 12, 30) + datetime.timedelta(days=int(serial_m.group(1)))
            cur = {'year': dt.year, 'month': dt.month, 'day': dt.day,
                   'dow': _DOW[dt.weekday()], 'lunch': [], 'mats': [], 'snack': ''}
            day_entries.append(cur)
        elif cur is None or not (name_v or mats_v or snack_cell):
            continue

        if name_v:
            cur['lunch'].append(name_v)
        cur['mats'].extend(mats_v)
        if snack_cell and not cur['snack']:
            snack_lines = snack_cell.split('\n')
            cur['snack'] = snack_lines[0].strip()
            if len(snack_lines) > 1:
                snack_ing = ' '.join(snack_lines[1:])
                snack_ing = snack_ing.replace('（', '').replace('）', '').strip()
                if snack_ing:
                    cur['mats'].append(snack_ing)

    if not day_entries:
        return ""

    y0, mo0 = day_entries[0]['year'], day_entries[0]['month']
    lines = [f"# 献立データ {y0}年{mo0:02d}月", ""]
    for e in day_entries:
        lines.append(f"【{e['month']}/{e['day']}({e['dow']})】")
        if e['lunch']:
            lines.append(f"昼食: {' / '.join(e['lunch'])}")
        if e['snack']:
            lines.append(f"おやつ: {e['snack']}")
        if e['mats']:
            lines.append(f"材料: {', '.join(e['mats'])}")
        lines.append("")

    return "\n".join(lines)


def _excel_to_text_iwato(df, fname=""):
    """岩戸こども園様形式（離乳食・後期/完了期の2段おやつ）→ 構造化テキスト。
    タイトル行や園名の記載がなく、ヘッダー行の「献立名」「材料」の位置から
    判定する。1日2行構成（1行目：日付＋主菜＋材料（最大4列）＋調味料3列＋
    後期/完了期おやつ2列、2行目：副菜＋材料継続）。おやつ欄はセル内改行で
    「おやつ名\\n（材料）」の形式になっているため、名前と材料を分離して扱う。"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    header_row = None
    for r in range(n_rows):
        if cv(r, 2) == '献立名' and '材料' in cv(r, 3):
            header_row = r
            break
    if header_row is None:
        return ""

    snack_cols = [c for c in range(4, n_cols) if any(k in cv(header_row, c) for k in ('後期', '完了期'))]
    if not snack_cols:
        snack_cols = [c for c in (10, 11) if c < n_cols]
    mat_cols = [c for c in range(3, n_cols) if c not in snack_cols]

    _DOW = ['月', '火', '水', '木', '金', '土', '日']
    _SERIAL_RE = re.compile(r'^(\d{5})(\.0)?$')

    def _split_snack(cell):
        lines = cell.split('\n')
        name = lines[0].strip()
        ing_text = ''
        if len(lines) > 1:
            ing_text = re.sub(r'[（(）)]', '', ' '.join(lines[1:])).strip()
        return name, ing_text

    day_entries = []
    cur = None
    for r in range(header_row + 1, n_rows):
        col0 = cv(r, 0)
        if col0.startswith(('※', '＊印', '*印')):
            break

        ymd_m = re.match(r'(\d{4})-(\d{2})-(\d{2})', col0)
        serial_m = None if ymd_m else _SERIAL_RE.match(col0)
        name_v = cv(r, 2)
        mats_v = [cv(r, c) for c in mat_cols if cv(r, c)]
        snack_cells = [cv(r, c) for c in snack_cols if cv(r, c)]

        if ymd_m:
            y, mo, d = int(ymd_m.group(1)), int(ymd_m.group(2)), int(ymd_m.group(3))
            cur = {'year': y, 'month': mo, 'day': d,
                   'dow': _DOW[datetime.date(y, mo, d).weekday()],
                   'lunch': [], 'mats': [], 'snack': []}
            day_entries.append(cur)
        elif serial_m:
            dt = datetime.date(1899, 12, 30) + datetime.timedelta(days=int(serial_m.group(1)))
            cur = {'year': dt.year, 'month': dt.month, 'day': dt.day,
                   'dow': _DOW[dt.weekday()], 'lunch': [], 'mats': [], 'snack': []}
            day_entries.append(cur)
        elif cur is None or not (name_v or mats_v or snack_cells):
            continue

        if name_v:
            cur['lunch'].append(name_v)
        cur['mats'].extend(mats_v)
        for sc in snack_cells:
            sname, sing = _split_snack(sc)
            if sname and sname not in cur['snack']:
                cur['snack'].append(sname)
            if sing:
                cur['mats'].append(sing)

    if not day_entries:
        return ""

    y0, mo0 = day_entries[0]['year'], day_entries[0]['month']
    lines = [f"# 献立データ {y0}年{mo0:02d}月", ""]
    for e in day_entries:
        lines.append(f"【{e['month']}/{e['day']}({e['dow']})】")
        if e['lunch']:
            lines.append(f"昼食: {' / '.join(e['lunch'])}")
        if e['snack']:
            lines.append(f"おやつ: {' / '.join(e['snack'])}")
        if e['mats']:
            lines.append(f"材料: {', '.join(e['mats'])}")
        lines.append("")

    return "\n".join(lines)


def _excel_to_text_kunimi(df):
    """
    くにみ子ども園形式（横並び・4列/日・午前/昼食/午後の3食区分・乳児幼児2分量）
    → 構造化テキスト。
    列構成（date_col=日付「N日」があるセル）：
      献立名（料理名）: date_col - 1
      材料名          : date_col
      乳児用グラム    : date_col + 1
      幼児用グラム    : date_col + 2
    「午前」「昼食」「午後」の各ラベル行（col1）の直後に「乳児/幼児」ヘッダー行が
    現れ、そこから次のラベル行までが材料表（1行1食材）になっている。
    """
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    # タイトルの年月表記がシート右側の遠い列にあるため、全列を対象に広く探索する
    year_month, month_num = "", 0
    for r in range(min(5, n_rows)):
        for c in range(n_cols):
            m = re.match(r'(\d{4})年(\d{1,2})月', cv(r, c))
            if m:
                month_num = int(m.group(2))
                year_month = f"{m.group(1)}年{month_num:02d}月"
                break
        if year_month:
            break

    # 日付行検出：「N日」パターンが3個以上ある行
    date_row, date_cols = None, {}
    for r in range(n_rows):
        temp = {c: cv(r, c) for c in range(n_cols) if re.match(r'^\d+日$', cv(r, c))}
        if len(temp) >= 3:
            date_row, date_cols = r, temp
            break
    if date_row is None:
        return ""

    def date_label(col_c):
        dm = re.match(r'(\d+)日', date_cols[col_c])
        dow = cv(date_row, col_c + 1).strip('()（）')
        if dm and month_num:
            return f"{month_num}/{dm.group(1)}({dow})" if dow else f"{month_num}/{dm.group(1)}"
        return date_cols[col_c]

    # 食事区分（午前・昼食・午後）ラベル行と、直後の乳児/幼児材料ヘッダー行を検出
    MEAL_LABELS = ('午前', '昼食', '午後')
    header_rows = [r for r in range(date_row + 1, n_rows)
                   if cv(r, 4) == '乳児' and cv(r, 5) == '幼児']
    # col1にラベルがある行（食事区分ラベルだけでなく、末尾の栄養価サマリー行
    # 「脂肪」「カルシウム」等も含む）を材料表の終端検出に使う
    labeled_rows = sorted(r for r in range(date_row + 1, n_rows) if cv(r, 1))

    meal_blocks = []  # [(label, dish_start_row, mat_header_row), ...]
    for r in range(date_row + 1, n_rows):
        label = cv(r, 1)
        if label in MEAL_LABELS:
            nxt_header = next((h for h in header_rows if h > r), None)
            if nxt_header is not None:
                meal_blocks.append((label, r, nxt_header))

    lines = []
    if year_month:
        lines += [f"# 献立データ {year_month}", ""]

    for col_c in sorted(date_cols.keys()):
        lines.append(f"【{date_label(col_c)}】")
        lunch_dishes, snack_am, snack_pm = [], [], []
        all_mats = []

        for idx, (label, dish_row, mat_row) in enumerate(meal_blocks):
            block_end = next((lr for lr in labeled_rows if lr > mat_row), n_rows)

            if label == '昼食':
                for r in range(dish_row, mat_row):
                    v = cv(r, col_c - 1)
                    if v:
                        lunch_dishes.append(v)

            mats = []
            for r in range(mat_row + 1, block_end):
                v = cv(r, col_c)
                if v:
                    mats.append(v)
            all_mats.extend(mats)
            if label == '午前':
                snack_am.extend(mats)
            elif label == '午後':
                snack_pm.extend(mats)

        if lunch_dishes:
            lines.append(f"昼食: {' / '.join(lunch_dishes)}")
        if snack_am or snack_pm:
            lines.append(f"おやつ: {' / '.join(snack_am + snack_pm)}")
            if snack_am:
                lines.append(f"午前おやつ: {' / '.join(snack_am)}")
            if snack_pm:
                lines.append(f"午後おやつ: {' / '.join(snack_pm)}")
        if all_mats:
            lines.append(f"材料: {', '.join(all_mats)}")
        lines.append("")

    return '\n'.join(lines)


def get_sheet_names(uploaded_file):
    engine = "xlrd" if uploaded_file.name.lower().endswith(".xls") else "openpyxl"
    try:
        xl = pd.ExcelFile(uploaded_file, engine=engine)
        return xl.sheet_names
    except Exception as e:
        st.error(f"ファイルの読み込みに失敗しました: {e}")
        return []


def excel_to_text(uploaded_file, sheet_name):
    engine = "xlrd" if uploaded_file.name.lower().endswith(".xls") else "openpyxl"
    df = pd.read_excel(
        uploaded_file, sheet_name=sheet_name, header=None, dtype=str, engine=engine
    )

    # フォーマット自動検出 → 専用パーサーに振り分け
    fmt = _detect_sheet_format(df)
    if fmt == 'sakae':
        return _excel_to_text_sakae(df, uploaded_file.name)
    if fmt == 'omiya':
        return _excel_to_text_omiya(df, uploaded_file.name)
    if fmt == 'mebaenomori':
        return _excel_to_text_mebaenomori(df)
    if fmt == 'yumehana':
        return _excel_to_text_yumehana(df, uploaded_file.name)
    if fmt == 'yamazaki':
        return _excel_to_text_yamazaki(df, uploaded_file.name)
    if fmt == 'ayumi':
        return _excel_to_text_ayumi(df)
    if fmt == 'kyomachibori':
        return _excel_to_text_kyomachibori(df)
    if fmt == 'tomikiya':
        return _excel_to_text_tomikiya(df)
    if fmt == 'sakae_baby':
        return _excel_to_text_sakae_baby(df)
    if fmt == 'miyama':
        return _excel_to_text_miyama(df)
    if fmt == 'kunimi':
        return _excel_to_text_kunimi(df)
    if fmt == 'kitanoda':
        return _excel_to_text_kitanoda(df, uploaded_file.name)
    if fmt == 'kitanoda_baby':
        return _excel_to_text_kitanoda_baby(df, uploaded_file.name)
    if fmt == 'iwato':
        return _excel_to_text_iwato(df, uploaded_file.name)

    n_rows, n_cols = df.shape

    def cell_val(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    # 日付行を全行から検索（「N日(曜)」「N月M日(曜)」パターンが3個以上ある行をブロック開始とする）
    _WIDE_DATE_RE = re.compile(r'^(?:(\d{1,2})月)?(\d{1,2})日\(([月火水木金土日])\)$')
    blocks = []  # list of (date_row_idx, {col: date_str})
    for r in range(n_rows):
        temp = {}
        for c in range(n_cols):
            v = cell_val(r, c)
            if _WIDE_DATE_RE.match(v):
                temp[c] = v
        if len(temp) >= 3:
            blocks.append((r, temp))

    # 日付構造が見つからない場合は行ごと出力にフォールバック
    if not blocks:
        rows = []
        for i, row in df.iterrows():
            cells = [
                str(v).strip() if pd.notna(v) and str(v).strip() not in ("nan", "") else ""
                for v in row
            ]
            if any(cells):
                rows.append(f"行{i + 1}: " + " | ".join(cells))
        return "\n".join(rows)

    # 年月を抽出（例：「2026年09月」）。タイトル・ファイル名の全角数字や
    # 「N月のみ」表記、年の記載なしにも対応する共通ヘルパーを使う。
    _, _month_num_i, year_month = _extract_year_month(df, uploaded_file.name)
    month_num = str(_month_num_i) if _month_num_i else ""
    # 日付ラベルに月が明記されている列（例：「8月31日」）が出た時点でこの値を
    # 更新し、以降の「N日」のみの列（月をまたいでも再掲されない）に引き継ぐ。
    cur_month = month_num

    skip_vals = {"[昼]", "[午後]", "献立名", "材料", "日付"}

    def is_valid_cell(v):
        if not v or v in skip_vals:
            return False
        if v.startswith("※"):  # 注記・免責文を除外
            return False
        return True

    lines = []
    if year_month:
        lines.append(f"# 献立データ {year_month}")
        lines.append("")

    for block_idx, (block_date_row, block_date_cols) in enumerate(blocks):
        # ブロック終端（次のブロック開始行 or ファイル末尾）
        block_end = blocks[block_idx + 1][0] if block_idx + 1 < len(blocks) else n_rows

        # このブロック内の材料セクション開始行
        block_mat_row = None
        for r in range(block_date_row + 1, block_end):
            for c in range(min(5, n_cols)):
                if cell_val(r, c) == "材料":
                    block_mat_row = r
                    break
            if block_mat_row is not None:
                break

        dish_end = block_mat_row if block_mat_row is not None else block_end

        for col_c in sorted(block_date_cols.keys()):
            raw_date = block_date_cols[col_c]
            dm = _WIDE_DATE_RE.match(raw_date)
            if dm.group(1):
                cur_month = dm.group(1)
            if cur_month:
                date_label = f"{cur_month}/{dm.group(2)}({dm.group(3)})"
            else:
                date_label = raw_date

            lines.append(f"【{date_label}】")

            # [午後]マーカーの行を検索（その日付列の1列前に出現する）
            afternoon_start = dish_end
            if col_c > 0:
                for r in range(block_date_row + 1, dish_end):
                    if cell_val(r, col_c - 1) == "[午後]":
                        afternoon_start = r
                        break

            # 昼食献立
            lunch = []
            for r in range(block_date_row + 1, afternoon_start):
                v = cell_val(r, col_c)
                if is_valid_cell(v):
                    lunch.append(v)

            # おやつ
            snack = []
            for r in range(afternoon_start, dish_end):
                v = cell_val(r, col_c)
                if is_valid_cell(v):
                    snack.append(v)

            # 材料（ブロック内のみ）
            mats = []
            if block_mat_row is not None:
                for r in range(block_mat_row, block_end):
                    v = cell_val(r, col_c)
                    if is_valid_cell(v):
                        mats.append(v)

            if lunch:
                lines.append(f"昼食: {' / '.join(lunch)}")
            if snack:
                lines.append(f"おやつ: {' / '.join(snack)}")
            if mats:
                lines.append(f"材料: {', '.join(mats)}")
            lines.append("")

    return "\n".join(lines)


def table_to_docx(markdown_text, title=None):
    from docx.oxml.ns import qn as _qn

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # 用紙サイズを A4縦 として明示
    pgSz = section._sectPr.find(_qn('w:pgSz'))
    if pgSz is not None:
        pgSz.set(_qn('w:orient'), 'portrait')
        pgSz.set(_qn('w:code'), '9')

    # ページヘッダーにタイトルを設定（Wordの機能により2枚目以降にも自動で表示される）
    if title:
        header_para = section.header.paragraphs[0]
        run = header_para.add_run(title)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    lines = markdown_text.strip().split('\n')
    table_lines = [
        l for l in lines
        if l.strip().startswith('|') and not re.match(r'\|[\s\-|]+\|', l.strip())
    ]

    if not table_lines:
        doc.add_paragraph(markdown_text)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
    rows = [[c.strip() for c in line.split('|')[1:-1]] for line in table_lines[1:]]
    n_cols = len(headers)
    rows = [r[:n_cols] + [''] * max(0, n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.allow_autofit = False

    # A4縦 本文幅 18cm を 4列に分配（日付 / 献立名 / おやつ / 結果）
    COL_WIDTHS = [Cm(2.0), Cm(5.5), Cm(3.5), Cm(7.0)]
    if n_cols != 4:
        each = Cm(18.0 / n_cols)
        COL_WIDTHS = [each] * n_cols

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.width = COL_WIDTHS[i]
        para = cell.paragraphs[0]
        if para.runs:
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(9)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            cell.width = COL_WIDTHS[c_idx]
            para = cell.paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.size = Pt(8)
                if c_idx == n_cols - 1 and val.startswith('●'):
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


_XLS_HALIGN = {1: 'left', 2: 'center', 3: 'right', 4: 'fill', 5: 'justify',
               6: 'centerContinuous', 7: 'distributed'}
_XLS_VALIGN = {0: 'top', 1: 'center', 2: 'bottom', 3: 'justify', 4: 'distributed'}


def _xls_font_to_openpyxl(xls_book, font_index):
    """xlrdのフォント情報をopenpyxlのFontに変換（フォント名・サイズ・太字・斜体・色）。"""
    try:
        f = xls_book.font_list[font_index]
    except (IndexError, TypeError):
        return None
    color = None
    if f.colour_index and f.colour_index != 32767:
        rgb = xls_book.colour_map.get(f.colour_index)
        if rgb:
            color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    return Font(
        name=f.name or None,
        size=(f.height / 20) if f.height else None,
        bold=bool(f.bold),
        italic=bool(f.italic),
        color=color,
    )


def _xls_alignment_to_openpyxl(xf):
    """xlrdのXF配置情報をopenpyxlのAlignmentに変換。"""
    return Alignment(
        horizontal=_XLS_HALIGN.get(xf.alignment.hor_align),
        vertical=_XLS_VALIGN.get(xf.alignment.vert_align),
        wrap_text=bool(xf.alignment.text_wrapped),
    )


def _xls_print_area_range(xls_book, sheet_index):
    """xlrdの定義済み名前「Print_Area」から該当シートの印刷範囲を
    openpyxlの print_area 文字列（例："A1:BN105"）として返す。無ければNone。"""
    names = xls_book.name_map.get('print_area')
    if not names:
        return None
    for name_obj in names:
        try:
            sheet_obj, row_lo, row_hi, col_lo, col_hi = name_obj.area2d()
        except Exception:
            continue
        if sheet_obj.number != sheet_index:
            continue
        if row_hi <= row_lo or col_hi <= col_lo:
            continue
        start = f'{get_column_letter(col_lo + 1)}{row_lo + 1}'
        end = f'{get_column_letter(col_hi)}{row_hi}'
        return f'{start}:{end}'
    return None


def _apply_a4_landscape_fit(ws):
    """印刷設定をA4横・1ページに収まるよう明示的に設定する。
    xlrdは用紙の向き・拡大縮小率（SETUPレコード）を読み取れないため、元ファイルの
    設定を複製することはできない。その代わりA4横向き・縦横1ページに固定することで、
    献立表（横に日付が並ぶ横長レイアウト）が常にきれいに1ページへ収まるようにする。"""
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True


def _apply_kunimi_print_layout(ws):
    """くにみ子ども園形式限定：印刷用にA4横1枚へ収まるようレイアウトを調整する。
    「エネルギー」行（栄養価サマリーの先頭）の位置は献立内容によって月ごとに
    変わるため、固定行番号は使わずすべて実データから動的に検出する。
    くにみ形式でない（日付行や午前/昼食/午後/エネルギーの見出しが見当たらない）
    場合は何もしない（安全側）。"""
    max_row, max_col = ws.max_row, ws.max_column

    def cv(r, c):
        v = ws.cell(row=r, column=c).value
        v = "" if v is None else str(v).strip()
        return "" if v in ("nan", "None") else v

    # 日付行検出（「N日」パターンが3個以上ある行）
    date_row, date_cols = None, []
    for r in range(1, max_row + 1):
        cols = [c for c in range(1, max_col + 1) if re.match(r'^\d+日$', cv(r, c))]
        if len(cols) >= 3:
            date_row, date_cols = r, cols
            break
    if date_row is None:
        return

    labeled_rows = sorted(r for r in range(date_row + 1, max_row + 1) if cv(r, 2))
    meal_label_rows = [r for r in labeled_rows if cv(r, 2) in ('午前', '昼食', '午後')]
    energy_row = next((r for r in labeled_rows if cv(r, 2) == 'エネルギー'), None)
    if not meal_label_rows or energy_row is None:
        return

    # 1. 乳児/幼児グラム列（日付列の直後2列）の幅を広げる（###表示防止）
    for c in date_cols:
        for gc in (c + 1, c + 2):
            if gc <= max_col:
                ws.column_dimensions[get_column_letter(gc)].width = 5.5

    # 2. 区分ラベル（午前/昼食/午後）セルを改行入り・中央揃え・折り返し表示に
    #    （既存の縦結合はそのまま。幅が狭い列でも1文字ずつ改行されて読めるようにする）
    for label_row in meal_label_rows:
        label = cv(label_row, 2)
        cell = ws.cell(row=label_row, column=2)
        cell.value = '\n'.join(label)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # 3. 空行を非表示。ただし複数列にまたがる結合セルが1つでもある行は、その月の
    #    献立内容によっては空欄なだけの「予備の材料行」である可能性があるため隠さない
    #    （区分ラベル列＝2列目だけの縦結合は対象外。ラベルのある行に文字があるため
    #    自然に保護されており、そこだけを理由に他の行を保護する必要はない）。
    protected_rows = set()
    for merged_range in ws.merged_cells.ranges:
        if merged_range.max_col - merged_range.min_col < 1:
            continue  # 1列だけの結合（区分ラベル列の縦結合）は対象外
        protected_rows.update(range(merged_range.min_row, merged_range.max_row + 1))
    for r in range(date_row + 1, energy_row):
        if r in protected_rows:
            continue
        if all(cv(r, c) == "" for c in range(1, max_col + 1)):
            ws.row_dimensions[r].hidden = True

    # 4. 印刷範囲を「A1〜エネルギー行」に設定（栄養素の内訳行は印刷対象外）
    ws.print_area = f'A1:{get_column_letter(max_col)}{energy_row}'

    # 5. A4横1枚に上下左右の余白なく収まるよう、固定の縮小率＋行高さのスケーリングで
    #    調整する（fitToWidth/fitToHeightの自動計算に任せると、幅基準の縮小率が
    #    高さに対しては控えめすぎて上下に不自然な余白が残ることがあるため、
    #    幅基準の縮小率を先に求め、その縮小率で高さもちょうど1ページに収まるよう
    #    行の高さ自体を伸縮させる）。
    # 幅方向にどれだけ縮小が必要かは、Excelの列幅（文字単位）→ポイント換算に
    # 近似誤差が伴い、自前で縮小率を決め打ちすると横方向がはみ出す恐れがある
    # （実測で発生）。そのため縮小率自体はExcelの「幅・高さを1ページに収める」
    # 機能（fitToWidth/fitToHeight、_apply_a4_landscape_fitで設定済み）に委ね、
    # 確実にはみ出さないようにする。その上で、行の高さを「その概算縮小率で
    # 高さもほぼ1ページ分になる」比率に事前調整しておくことで、幅基準と高さ基準の
    # 縮小率の差を縮め、上下（または左右）に不自然な余白が残る問題を緩和する。
    MARGIN_PT = 14.0
    PAGE_W_PT, PAGE_H_PT = 841.89, 595.28  # A4横（297mm×210mm）をポイント換算
    printable_w = PAGE_W_PT - 2 * MARGIN_PT
    printable_h = PAGE_H_PT - 2 * MARGIN_PT
    PT_PER_CHAR_UNIT = 5.0  # 列幅（文字単位）→ポイントの近似換算係数（あくまで目安）

    total_w_units = sum(
        (ws.column_dimensions[get_column_letter(c)].width
         if get_column_letter(c) in ws.column_dimensions
         and ws.column_dimensions[get_column_letter(c)].width else 8.43)
        for c in range(1, max_col + 1)
    )
    total_w_pt_est = total_w_units * PT_PER_CHAR_UNIT

    total_h_pt = sum(
        (ws.row_dimensions[r].height if ws.row_dimensions[r].height else 15.0)
        for r in range(1, energy_row + 1)
        if not ws.row_dimensions[r].hidden
    )
    if total_w_pt_est > 0 and total_h_pt > 0:
        target_h_pt = total_w_pt_est * (printable_h / printable_w)
        stretch = target_h_pt / total_h_pt
        for r in range(1, energy_row + 1):
            if ws.row_dimensions[r].hidden:
                continue
            cur_h = ws.row_dimensions[r].height if ws.row_dimensions[r].height else 15.0
            ws.row_dimensions[r].height = cur_h * stretch

    for margin_attr in ('left', 'right', 'top', 'bottom'):
        setattr(ws.page_margins, margin_attr, MARGIN_PT / 72)
    ws.page_margins.header = 0
    ws.page_margins.footer = 0


def create_colored_excel(uploaded_file, color_groups=None):
    """
    全シートを色付きにして .xlsx で返す。フォーマット自動検出により
    さかえ・おおみや・ゆめのはな・既存形式それぞれの材料列に色付け。
    .xls はデータ・マージセルを再構築（書式は一部失われる）。
    color_groups: [{"keywords": [str,...], "color": "RRGGBB"}, ...]
                  省略時は DEFAULT_COLOR_GROUPS（従来の固定配色）を使用。
    """
    is_xls = uploaded_file.name.lower().endswith(".xls")
    file_bytes = uploaded_file.read()

    # ─── 色定義 ──────────────────────────────────────────────
    def _fill(hex6):
        return PatternFill(fill_type='solid', fgColor=hex6)

    ING_COLOR_RULES = [
        (g["keywords"], _fill(g["color"]))
        for g in (color_groups or DEFAULT_COLOR_GROUPS)
    ]

    # ─── フォーマット別の色付けロジック ──────────────────────
    def _apply_colors(cv_fn, af_fn, n_rows, n_cols, fmt):
        """cv_fn(r,c)→str、af_fn(r,c,fill)→None を使って色付け"""
        if fmt == 'sakae':
            # 列3-6が材料
            for r in range(n_rows):
                for c in range(3, min(7, n_cols)):
                    v = cv_fn(r, c)
                    if not v:
                        continue
                    for kw_list, fc in ING_COLOR_RULES:
                        if any(kw in v for kw in kw_list):
                            af_fn(r, c, fc)
                            break
        elif fmt == 'omiya':
            # 列2が全材料（1セル）
            for r in range(n_rows):
                v = cv_fn(r, 2)
                if not v:
                    continue
                for kw_list, fc in ING_COLOR_RULES:
                    if any(kw in v for kw in kw_list):
                        af_fn(r, 2, fc)
                        break
        elif fmt in ('yumehana', 'mebaenomori'):
            # 列3-9が材料（ゆめのはな・めばえの森共通）
            for r in range(n_rows):
                for c in range(3, min(10, n_cols)):
                    v = cv_fn(r, c)
                    if not v:
                        continue
                    for kw_list, fc in ING_COLOR_RULES:
                        if any(kw in v for kw in kw_list):
                            af_fn(r, c, fc)
                            break
        elif fmt == 'yamazaki':
            # 「N日(曜)」のcol_c+1が材料列（日付ごとに3列ずつ）
            date_cols = set()
            for r in range(n_rows):
                for c in range(n_cols):
                    v = cv_fn(r, c)
                    if re.match(r'^\d+日\([月火水木金土日]\)$', v):
                        date_cols.add(c)
            for col_c in date_cols:
                for r in range(n_rows):
                    v = cv_fn(r, col_c + 1)
                    if not v:
                        continue
                    for kw_list, fc in ING_COLOR_RULES:
                        if any(kw in v for kw in kw_list):
                            af_fn(r, col_c + 1, fc)
                            break
        elif fmt in ('ayumi', 'miyama'):
            # 歩学園・美山保育園: col3〜col9が材料
            for r in range(n_rows):
                for c in range(3, min(10, n_cols)):
                    v = cv_fn(r, c)
                    if not v:
                        continue
                    for kw_list, fc in ING_COLOR_RULES:
                        if any(kw in v for kw in kw_list):
                            af_fn(r, c, fc)
                            break
        else:
            # 既存形式: 「N日(曜)」横並びブロック検出
            blocks = []
            for r in range(n_rows):
                temp = {}
                for c in range(n_cols):
                    v = cv_fn(r, c)
                    if re.match(r'^\d+日\([月火水木金土日]\)$', v):
                        temp[c] = v
                if len(temp) >= 3:
                    blocks.append((r, temp))
            for b_idx, (b_row, b_cols) in enumerate(blocks):
                b_end = blocks[b_idx + 1][0] if b_idx + 1 < len(blocks) else n_rows
                mat_row = None
                for r in range(b_row + 1, b_end):
                    for c in range(min(5, n_cols)):
                        if cv_fn(r, c) == "材料":
                            mat_row = r
                            break
                    if mat_row is not None:
                        break
                if mat_row is None:
                    continue
                for col_c in sorted(b_cols.keys()):
                    for r in range(mat_row, b_end):
                        v = cv_fn(r, col_c)
                        if not v:
                            continue
                        for kw_list, fc in ING_COLOR_RULES:
                            if any(kw in v for kw in kw_list):
                                af_fn(r, col_c, fc)
                                break

    # ─── .xlsx ───────────────────────────────────────────────
    if not is_xls:
        wb = load_workbook(BytesIO(file_bytes), data_only=True)
        for sh_name in wb.sheetnames:
            ws = wb[sh_name]
            n_rows, n_cols = ws.max_row, ws.max_column
            df_sh = pd.read_excel(BytesIO(file_bytes), sheet_name=sh_name,
                                  header=None, dtype=str, engine='openpyxl')
            fmt = _detect_sheet_format(df_sh)

            def _cv(r, c, _ws=ws):
                if r < 0 or r >= _ws.max_row or c < 0 or c >= _ws.max_column:
                    return ""
                v = _ws.cell(row=r + 1, column=c + 1).value
                return "" if v is None else ("" if str(v).strip() in ("nan", "", "None") else str(v).strip())

            def _af(r, c, fill, _ws=ws):
                try:
                    _ws.cell(row=r + 1, column=c + 1).fill = fill
                except AttributeError:
                    pass

            _apply_colors(_cv, _af, n_rows, n_cols, fmt)

    # ─── .xls ────────────────────────────────────────────────
    else:
        import xlrd as _xlrd
        # formatting_info=True で列幅・行高さを取得
        xls_wb = _xlrd.open_workbook(file_contents=file_bytes, formatting_info=True)
        wb = Workbook()
        first = True
        for sh_name in xls_wb.sheet_names():
            xls_ws = xls_wb.sheet_by_name(sh_name)
            n_cols = xls_ws.ncols

            # データがある実際の最終行を特定（空行の大量コピーを防ぐ）
            last_data_row = 0
            for r_i in range(xls_ws.nrows):
                if any(str(xls_ws.cell_value(r_i, c_i)).strip() not in ('', 'nan', 'None')
                       for c_i in range(n_cols)):
                    last_data_row = r_i
            n_rows = last_data_row + 1

            if first:
                ws = wb.active
                ws.title = sh_name[:31]
                first = False
            else:
                ws = wb.create_sheet(title=sh_name[:31])

            # 値をコピー（数値・日付セルは型を保持し、文字列化しない）
            for r_i in range(n_rows):
                for c_i in range(n_cols):
                    v = xls_ws.cell_value(r_i, c_i)
                    if v is None or str(v).strip() in ("nan", "", "None"):
                        continue
                    if xls_ws.cell_type(r_i, c_i) == _xlrd.XL_CELL_TEXT:
                        ws.cell(row=r_i + 1, column=c_i + 1, value=str(v).strip() or None)
                    else:
                        ws.cell(row=r_i + 1, column=c_i + 1, value=v)

            # マージセルをコピー（実データ範囲内のみ）
            for row_lo, row_hi, col_lo, col_hi in xls_ws.merged_cells:
                if row_lo >= n_rows:
                    continue
                try:
                    ws.merge_cells(start_row=row_lo + 1, start_column=col_lo + 1,
                                   end_row=min(row_hi, n_rows), end_column=col_hi)
                except Exception:
                    pass

            # 列幅・行高さ（formatting_info=True で取得済み）
            try:
                for c_i in range(n_cols):
                    ci = xls_ws.colinfo_map.get(c_i)
                    if ci and ci.width > 0:
                        ws.column_dimensions[get_column_letter(c_i + 1)].width = ci.width / 256
            except Exception:
                pass
            try:
                for r_i in range(n_rows):
                    ri = xls_ws.rowinfo_map.get(r_i)
                    if ri and ri.height > 0:
                        ws.row_dimensions[r_i + 1].height = ri.height / 20
            except Exception:
                pass

            # 罫線をコピー
            _XLS_LINE = {
                1: 'thin', 2: 'medium', 3: 'dashed', 4: 'dotted',
                5: 'thick', 6: 'double', 7: 'hair', 8: 'mediumDashed',
                9: 'dashDot', 10: 'mediumDashDot', 11: 'dashDotDot',
                12: 'mediumDashDotDot', 13: 'slantDashDot',
            }

            def _side(line_type, colour_idx):
                style = _XLS_LINE.get(line_type)
                if not style:
                    return Side(border_style=None)
                rgb = xls_wb.colour_map.get(colour_idx)
                color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}' if rgb else '000000'
                return Side(border_style=style, color=color)

            for r_i in range(n_rows):
                for c_i in range(n_cols):
                    try:
                        xf = xls_wb.xf_list[xls_ws.cell_xf_index(r_i, c_i)]
                        b = xf.border
                        cell = ws.cell(row=r_i + 1, column=c_i + 1)
                        cell.border = Border(
                            left=_side(b.left_line_style, b.left_colour_index),
                            right=_side(b.right_line_style, b.right_colour_index),
                            top=_side(b.top_line_style, b.top_colour_index),
                            bottom=_side(b.bottom_line_style, b.bottom_colour_index),
                        )
                        font = _xls_font_to_openpyxl(xls_wb, xf.font_index)
                        if font is not None:
                            cell.font = font
                        cell.alignment = _xls_alignment_to_openpyxl(xf)
                        fmt = xls_wb.format_map.get(xf.format_key)
                        if fmt and fmt.format_str:
                            cell.number_format = fmt.format_str
                    except Exception:
                        pass

            # 印刷範囲（Print_Area）を元ファイルと同じ範囲に設定
            print_range = _xls_print_area_range(xls_wb, xls_wb.sheet_names().index(sh_name))
            if print_range:
                ws.print_area = print_range
            _apply_a4_landscape_fit(ws)

            # フォーマット検出
            df_sh = pd.read_excel(BytesIO(file_bytes), sheet_name=sh_name,
                                  header=None, dtype=str, engine='xlrd')
            fmt = _detect_sheet_format(df_sh)

            def _cv(r, c, _xws=xls_ws):
                if r < 0 or r >= _xws.nrows or c < 0 or c >= _xws.ncols:
                    return ""
                v = _xws.cell_value(r, c)
                return "" if v is None else ("" if str(v).strip() in ("nan", "", "None") else str(v).strip())

            def _af(r, c, fill, _ws=ws):
                try:
                    _ws.cell(row=r + 1, column=c + 1).fill = fill
                except AttributeError:
                    pass

            _apply_colors(_cv, _af, n_rows, n_cols, fmt)

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio


def _kunimi_dish_context_map(df):
    """くにみ子ども園形式限定：昼食の材料セル(row, col)→その日の昼食献立名テキスト、
    の対応表を返す。条件付き置換ルール（例：「中華麺」は献立名がラーメンなら維持し、
    皿うどん/ちゃんぽん/焼きそばなら置換）が、そのセルがどの日のどの献立の材料かを
    判定するために使う。くにみ形式でなければ空dictを返す（＝条件付き置換は発火しない）。"""
    n_rows, n_cols = df.shape

    def cv(r, c):
        if r < 0 or r >= n_rows or c < 0 or c >= n_cols:
            return ""
        v = str(df.iloc[r, c]).strip()
        return "" if v in ("nan", "", "None") else v

    date_row, date_cols = None, {}
    for r in range(n_rows):
        temp = {c: cv(r, c) for c in range(n_cols) if re.match(r'^\d+日$', cv(r, c))}
        if len(temp) >= 3:
            date_row, date_cols = r, temp
            break
    if date_row is None:
        return {}

    header_rows = [r for r in range(date_row + 1, n_rows)
                   if cv(r, 4) == '乳児' and cv(r, 5) == '幼児']
    labeled_rows = sorted(r for r in range(date_row + 1, n_rows) if cv(r, 1))

    meal_blocks = []  # (dish_row, mat_row) の昼食ブロックのみ
    for r in range(date_row + 1, n_rows):
        if cv(r, 1) == '昼食':
            nxt_header = next((h for h in header_rows if h > r), None)
            if nxt_header is not None:
                meal_blocks.append((r, nxt_header))

    context = {}
    for col_c in date_cols:
        for dish_row, mat_row in meal_blocks:
            dish_text = ' '.join(
                cv(r, col_c - 1) for r in range(dish_row, mat_row) if cv(r, col_c - 1)
            )
            block_end = next((lr for lr in labeled_rows if lr > mat_row), n_rows)
            for r in range(mat_row + 1, block_end):
                context[(r, col_c)] = dish_text
    return context


def apply_replacements_to_excel(uploaded_file, pairs):
    """
    アップロードされたExcelの全セルに対し、pairs の (from → to) を順に文字列置換して
    新しい .xlsx を返す。列オフセット等を考慮しないフォーマット非依存の一律置換が基本だが、
    pairに"when_dish"/"unless_dish"（その日の昼食献立名に含む/含まないキーワードのリスト）
    が指定されている場合は、くにみ子ども園形式に限り献立名の文脈を見て条件判定する
    （それ以外の形式では条件付きpairは常にスキップ＝置換されない。無条件pairの動作は不変）。
    .xls はデータ・マージセル・列幅行高さ・罫線を再構築する（create_colored_excelと同様）。
    pairs: [{"from": str, "to": str, "when_dish": [str,...]?, "unless_dish": [str,...]?}, ...]
    """
    is_xls = uploaded_file.name.lower().endswith(".xls")
    file_bytes = uploaded_file.read()
    has_conditional = any(p.get("when_dish") or p.get("unless_dish") for p in pairs)

    def _replace(v, dish_text=""):
        s = str(v)
        for p in pairs:
            if not p.get("from"):
                continue
            when_dish = p.get("when_dish")
            unless_dish = p.get("unless_dish")
            if when_dish and not any(k in dish_text for k in when_dish):
                continue
            if unless_dish and any(k in dish_text for k in unless_dish):
                continue
            s = s.replace(p["from"], p.get("to", ""))
        return s

    def _sheet_format(sh_name, engine):
        try:
            df = pd.read_excel(BytesIO(file_bytes), sheet_name=sh_name,
                                header=None, dtype=str, engine=engine)
        except Exception:
            return None, None
        return _detect_sheet_format(df), df

    def _sheet_dish_context(fmt, df):
        if not has_conditional or fmt != 'kunimi' or df is None:
            return {}
        return _kunimi_dish_context_map(df)

    # ─── .xlsx ───────────────────────────────────────────────
    if not is_xls:
        wb = load_workbook(BytesIO(file_bytes))
        for sh_name in wb.sheetnames:
            ws = wb[sh_name]
            fmt, df = _sheet_format(sh_name, "openpyxl")
            dish_ctx = _sheet_dish_context(fmt, df)
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        ctx = dish_ctx.get((cell.row - 1, cell.column - 1), "")
                        cell.value = _replace(cell.value, ctx)
            if fmt == 'kunimi':
                _apply_a4_landscape_fit(ws)
                _apply_kunimi_print_layout(ws)
        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)
        return bio

    # ─── .xls ────────────────────────────────────────────────
    import xlrd as _xlrd
    xls_wb = _xlrd.open_workbook(file_contents=file_bytes, formatting_info=True)
    wb = Workbook()
    first = True
    for sh_name in xls_wb.sheet_names():
        xls_ws = xls_wb.sheet_by_name(sh_name)
        n_cols = xls_ws.ncols

        last_data_row = 0
        for r_i in range(xls_ws.nrows):
            if any(str(xls_ws.cell_value(r_i, c_i)).strip() not in ('', 'nan', 'None')
                   for c_i in range(n_cols)):
                last_data_row = r_i
        n_rows = last_data_row + 1

        if first:
            ws = wb.active
            ws.title = sh_name[:31]
            first = False
        else:
            ws = wb.create_sheet(title=sh_name[:31])

        fmt, _df = _sheet_format(sh_name, "xlrd")
        dish_ctx = _sheet_dish_context(fmt, _df)

        # 値をコピー（置換は文字列セルのみに適用。数値・日付セルは型を保持する）
        for r_i in range(n_rows):
            for c_i in range(n_cols):
                v = xls_ws.cell_value(r_i, c_i)
                if v is None or str(v).strip() in ("nan", "", "None"):
                    continue
                if xls_ws.cell_type(r_i, c_i) == _xlrd.XL_CELL_TEXT:
                    ctx = dish_ctx.get((r_i, c_i), "")
                    ws.cell(row=r_i + 1, column=c_i + 1, value=_replace(str(v).strip(), ctx) or None)
                else:
                    ws.cell(row=r_i + 1, column=c_i + 1, value=v)

        # マージセルをコピー
        for row_lo, row_hi, col_lo, col_hi in xls_ws.merged_cells:
            if row_lo >= n_rows:
                continue
            try:
                ws.merge_cells(start_row=row_lo + 1, start_column=col_lo + 1,
                               end_row=min(row_hi, n_rows), end_column=col_hi)
            except Exception:
                pass

        # 列幅・行高さ
        try:
            for c_i in range(n_cols):
                ci = xls_ws.colinfo_map.get(c_i)
                if ci and ci.width > 0:
                    ws.column_dimensions[get_column_letter(c_i + 1)].width = ci.width / 256
        except Exception:
            pass
        try:
            for r_i in range(n_rows):
                ri = xls_ws.rowinfo_map.get(r_i)
                if ri and ri.height > 0:
                    ws.row_dimensions[r_i + 1].height = ri.height / 20
        except Exception:
            pass

        # 罫線をコピー
        _XLS_LINE = {
            1: 'thin', 2: 'medium', 3: 'dashed', 4: 'dotted',
            5: 'thick', 6: 'double', 7: 'hair', 8: 'mediumDashed',
            9: 'dashDot', 10: 'mediumDashDot', 11: 'dashDotDot',
            12: 'mediumDashDotDot', 13: 'slantDashDot',
        }

        def _side(line_type, colour_idx):
            style = _XLS_LINE.get(line_type)
            if not style:
                return Side(border_style=None)
            rgb = xls_wb.colour_map.get(colour_idx)
            color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}' if rgb else '000000'
            return Side(border_style=style, color=color)

        for r_i in range(n_rows):
            for c_i in range(n_cols):
                try:
                    xf = xls_wb.xf_list[xls_ws.cell_xf_index(r_i, c_i)]
                    b = xf.border
                    cell = ws.cell(row=r_i + 1, column=c_i + 1)
                    cell.border = Border(
                        left=_side(b.left_line_style, b.left_colour_index),
                        right=_side(b.right_line_style, b.right_colour_index),
                        top=_side(b.top_line_style, b.top_colour_index),
                        bottom=_side(b.bottom_line_style, b.bottom_colour_index),
                    )
                    font = _xls_font_to_openpyxl(xls_wb, xf.font_index)
                    if font is not None:
                        cell.font = font
                    cell.alignment = _xls_alignment_to_openpyxl(xf)
                    numfmt = xls_wb.format_map.get(xf.format_key)
                    if numfmt and numfmt.format_str:
                        cell.number_format = numfmt.format_str
                except Exception:
                    pass

        # 印刷範囲（Print_Area）を元ファイルと同じ範囲に設定
        print_range = _xls_print_area_range(xls_wb, xls_wb.sheet_names().index(sh_name))
        if print_range:
            ws.print_area = print_range
        _apply_a4_landscape_fit(ws)
        if fmt == 'kunimi':
            _apply_kunimi_print_layout(ws)

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio


def compute_all_python_ngs(excel_text, rules_text="", leftover_words=None):
    """
    全ルールベースNGをPythonで計算。
    選択中ルール本文（rules_text）にステップ見出しが含まれるかで各チェックのON/OFFを判定する
    （ルールごとにチェック内容が異なるため、ハードコード項目を無条件に全ルールへ適用しない）。
    leftover_words: 置換ルールの「置換前」の単語リスト。指定時、置換されずに残っている単語を検出する
                    （事前に置換ルールを適用したはずのファイルで、置換漏れがないかの確認用）。
    Returns: (summary_text, day_ngs_dict)
      summary_text  … AIプロンプトに埋め込むNG一覧テキスト
      day_ngs_dict  … {date_str: [ng_str, ...]}
    """
    year, month_num, sorted_dates, entries = _parse_structured(excel_text)
    if not sorted_dates:
        return "", {}

    day_ngs = {ds: [] for ds in sorted_dates}

    def rule_has(*markers):
        return any(m in rules_text for m in markers)

    check_forbidden           = rule_has('使用禁止食材')
    check_consecutive_general = rule_has('食材の連続使用')
    check_cheese_consec       = rule_has('チーズ類の連続使用')
    check_meat3               = rule_has('肉類の3日連続使用', '肉類の３日連続使用')
    check_seasoning_consec    = rule_has('調味料・ベースの連続使用')
    check_imo                 = rule_has('芋類の連続・重複使用')
    check_same_day_dup        = rule_has('同じ日の材料内の重複')
    check_monthly_limit       = rule_has('月の使用上限')
    check_monday_holiday      = rule_has('月曜日・祝日翌日のチェック', '祝日翌日のチェック')
    check_weekly_tofu         = rule_has('週次 豆腐チェック', '週次豆腐チェック')
    check_weekly_menu         = rule_has('週次 献立チェック', '週次献立チェック')
    check_name_ingredient     = rule_has('献立名と材料の照合チェック')
    check_prep_ingredients    = rule_has('仕込み食材の集中チェック')
    check_monthly_count       = rule_has('献立名の出現回数チェック')
    check_saturday_noodle     = rule_has('土曜日」に麺', '土曜日に麺')
    check_jelly_prev          = rule_has('を含む文言がある日の前日')
    check_seasonal_stew       = rule_has('クリームシチュー')
    check_kunimi              = rule_has('くにみ子ども園固有ルール')

    def ing(ds):
        return entries[ds]['ingredients']

    def lunch(ds):
        return entries[ds]['lunch']

    def snack(ds):
        return entries[ds]['snack']

    def snack_am(ds):
        return entries[ds].get('snack_am', '')

    def snack_pm(ds):
        return entries[ds].get('snack_pm', '')

    def has(text, kw_list):
        return any(kw in text for kw in kw_list)

    # ── 置換漏れ検出（事前に置換ルールを適用したはずの単語が残っていないか） ──
    if leftover_words:
        for ds in sorted_dates:
            whole_text = ing(ds) + ' ' + lunch(ds) + ' ' + snack(ds)
            for w in leftover_words:
                if w and w in whole_text:
                    day_ngs[ds].append(f'● 置換漏れ：「{w}」が残っています（置換ルールの適用を確認してください）')

    # ── 禁止食材 ────────────────────────────────────────────────
    if check_forbidden:
        for ds in sorted_dates:
            for f in ['卵', 'マヨネーズ', '絹ごし豆腐']:
                if f in ing(ds):
                    day_ngs[ds].append(f'● {f}（使用禁止食材）')

    # ── チーズ2日連続 ──────────────────────────────────────────
    if check_cheese_consec:
        for i in range(1, len(sorted_dates)):
            if 'チーズ' in ing(sorted_dates[i]) and 'チーズ' in ing(sorted_dates[i-1]):
                day_ngs[sorted_dates[i]].append('● チーズ類2日連続')

    # ── 肉類3日連続 ────────────────────────────────────────────
    if check_meat3:
        for meat in MEAT_3DAY:
            for i in range(2, len(sorted_dates)):
                if all(meat in ing(sorted_dates[j]) for j in (i, i-1, i-2)):
                    day_ngs[sorted_dates[i]].append(f'● {meat}3日連続')

    # ── 調味料2日連続（日祝を挟めばOK） ──────────────────────
    if check_seasoning_consec:
        for seasoning in SEASONING_2DAY:
            for i in range(1, len(sorted_dates)):
                ds_c, ds_p = sorted_dates[i], sorted_dates[i-1]
                # 完全一致トークンで判定（部分一致だと「オイスターソース」等の
                # 複合語が汎用の「ソース」と誤って一致してしまうため）
                if seasoning in _split_ing(ing(ds_c)) and seasoning in _split_ing(ing(ds_p)):
                    d_c, d_p = _parse_date(ds_c, year), _parse_date(ds_p, year)
                    exempted = False
                    if d_c and d_p:
                        d = d_p + datetime.timedelta(days=1)
                        while d < d_c:
                            if d.weekday() == 6 or _is_holiday(d):
                                exempted = True
                                break
                            d += datetime.timedelta(days=1)
                    if not exempted:
                        day_ngs[ds_c].append(f'● {seasoning}2日連続')

    # ── 芋類3日連続 ────────────────────────────────────────────
    if check_imo:
        for i in range(2, len(sorted_dates)):
            if all(has(ing(sorted_dates[j]), IMO_KW) for j in (i, i-1, i-2)):
                day_ngs[sorted_dates[i]].append('● 芋類3日連続')

    # ── 汎用食材2日連続（免除リスト以外） ────────────────────
    if check_consecutive_general:
        _NUM_ONLY = re.compile(r'^\d+\.?\d*$')  # 数値トークン（量・シリアル等）を除外
        for i in range(1, len(sorted_dates)):
            ds_c, ds_p = sorted_dates[i], sorted_dates[i-1]
            toks_c = {t for t in _split_ing(ing(ds_c)) if not _is_exempt(t) and len(t) >= 2 and not _NUM_ONLY.match(t)}
            toks_p = {t for t in _split_ing(ing(ds_p)) if not _is_exempt(t) and len(t) >= 2 and not _NUM_ONLY.match(t)}
            for token in sorted(toks_c & toks_p):
                day_ngs[ds_c].append(f'● {token}2日連続')

    # ── 同日チェック ───────────────────────────────────────────
    if check_imo or check_same_day_dup:
        for ds in sorted_dates:
            i_text = ing(ds)
            toks = _split_ing(i_text)

            # 芋類2種類以上（同日）
            if check_imo:
                imo_cnt = sum(1 for kw in IMO_KW if kw in i_text)
                if imo_cnt >= 2:
                    day_ngs[ds].append(f'● 同日芋類{imo_cnt}種類（1種類のみ可）')

            if check_same_day_dup:
                # 酢2回以上
                if toks.count('酢') >= 2:
                    day_ngs[ds].append('● 同日「酢」2回使用')

                # みそ2回以上
                if toks.count('みそ') >= 2:
                    day_ngs[ds].append('● 同日「みそ」2回使用')

                # ほうれん草＋小松菜
                if 'ほうれん草' in i_text and '小松菜' in i_text:
                    day_ngs[ds].append('● 同日「ほうれん草」と「小松菜」重複')

                # 練り物2種類以上
                neri_cnt = sum(1 for kw in NERIMONO_KW if kw in i_text)
                if neri_cnt >= 2:
                    day_ngs[ds].append(f'● 同日練り物{neri_cnt}種類（1種類のみ可）')

                # 玉ねぎ3回以上
                tama_cnt = toks.count('玉ねぎ')
                if tama_cnt >= 3:
                    day_ngs[ds].append(f'● 同日「玉ねぎ」{tama_cnt}回使用')

                # 人参3回以上
                ninj_cnt = toks.count('人参')
                if ninj_cnt >= 3:
                    day_ngs[ds].append(f'● 同日「人参」{ninj_cnt}回使用')

    # ── 月上限（4回目に到達した日に記録） ─────────────────────
    if check_monthly_limit:
        for item in ['バター', 'チーズ', 'マヨドレ']:
            found = [ds for ds in sorted_dates if item in ing(ds)]
            if len(found) >= 4:
                day_ngs[found[3]].append(f'● {item}月4回目（上限超過）')

    # ── 月曜・祝日翌日チェック ────────────────────────────────
    if check_monday_holiday:
        for ds in sorted_dates:
            dow_i = _dow(ds)
            d = _parse_date(ds, year)
            is_mon = (dow_i == 0)
            prev_holiday = _is_holiday(d - datetime.timedelta(days=1)) if d else False
            is_mon_or_post = is_mon or prev_holiday
            is_thu = (dow_i == 3)
            i_text = ing(ds)
            ls_text = lunch(ds) + ' ' + snack(ds)

            if is_mon_or_post:
                for item in MON_NG_ITEMS:
                    if item in i_text:
                        day_ngs[ds].append(f'● 月曜/祝日翌日「{item}」NG')
                for m_kw in MUSHROOM_KW:
                    if m_kw in i_text:
                        day_ngs[ds].append(f'● 月曜/祝日翌日「{m_kw}」NG')
                        break
                for fish_kw in FISH_KW:   # ツナはOKなので FISH_KW（ツナなし）を使う
                    if fish_kw in i_text:
                        day_ngs[ds].append(f'● 月曜/祝日翌日「{fish_kw}」使用NG')
                        break
                # 月曜/祝日翌日に照り焼き・から揚げ・フライ
                for dish in ['照り焼き', 'から揚げ', 'フライ']:
                    if dish in ls_text:
                        check = ls_text.replace('フライドポテト', '') if dish == 'フライ' else ls_text
                        if dish in check:
                            day_ngs[ds].append(f'● 月曜/祝日翌日「{dish}」NG')

            # にら・青ねぎ：月曜・祝日翌日・木曜がNG
            if is_mon_or_post or is_thu:
                for nk in ['にら', '青ねぎ']:
                    if nk in i_text:
                        day_ngs[ds].append(f'● 月曜/祝日翌日/木曜「{nk}」NG')

    # ── 週次チェック（魚・豆腐・麺丼） ──────────────────────
    if check_weekly_tofu or check_weekly_menu:
        _dow_map = {'月': 0, '火': 1, '水': 2, '木': 3, '金': 4, '土': 5, '日': 6}

        def _week_key(ds):
            m2 = re.search(r'/(\d+)\(([月火水木金土日])\)', ds)
            return int(m2.group(1)) - _dow_map.get(m2.group(2), 0) if m2 else -1

        weeks = {}
        for ds in sorted_dates:
            weeks.setdefault(_week_key(ds), []).append(ds)

        for _, week_dates in sorted(weeks.items()):
            wd = sorted(week_dates, key=lambda d: _parse_date(d, year) or datetime.date.max)
            last = wd[-1]
            combined = ' '.join(lunch(ds) + ' ' + ing(ds) for ds in wd)

            if check_weekly_menu and not has(combined, FISH_WITH_TUNA):
                day_ngs[last].append('● 今週魚なし（週1回以上必要）')

            if check_weekly_tofu:
                tofu_ok = any(
                    has(ing(ds), TOFU_KW) or has(lunch(ds), TOFU_KW)
                    for ds in wd
                )
                if not tofu_ok:
                    day_ngs[last].append('● 今週豆腐なし（週1回以上必要）')

            if check_weekly_menu and not has(combined, NOODLE_KW):
                day_ngs[last].append('● 今週麺・丼なし（週1回以上必要）')

    # ── 仕込み食材7種類以上 ───────────────────────────────────
    if check_prep_ingredients:
        for ds in sorted_dates:
            i_text = ing(ds)
            found = [kw for kw in PREP_KW if kw in i_text]
            if len(found) >= 7:
                day_ngs[ds].append(f'● 仕込み食材{len(found)}種類（7種類以上）: {", ".join(found)}')

    # ── 月の献立回数上限 ──────────────────────────────────────
    if check_monthly_count:
        ingenge_days = [ds for ds in sorted_dates if 'いんげん' in (lunch(ds) + snack(ds))]
        fruit_days   = [ds for ds in sorted_dates if '果物' in (lunch(ds) + snack(ds))]
        if len(ingenge_days) >= 3:
            day_ngs[ingenge_days[2]].append('● 「いんげん」献立が月3回目（上限超過）')
        if len(fruit_days) >= 2:
            day_ngs[fruit_days[1]].append('● 「果物」献立が月2回目（上限超過）')

    # ── 土曜日に麺・丼なし ───────────────────────────────────
    if check_saturday_noodle:
        for ds in sorted_dates:
            if _dow(ds) == 5 and not has(lunch(ds), NOODLE_KW):
                day_ngs[ds].append('● 土曜日に麺・丼なし')

    # ── ハンバーグ/フライ前日おやつゼリーなし ────────────────
    if check_jelly_prev:
        for i in range(1, len(sorted_dates)):
            ds_today, ds_next = sorted_dates[i-1], sorted_dates[i]
            l_next = lunch(ds_next)
            need_jelly = 'ハンバーグ' in l_next
            if not need_jelly and 'フライ' in l_next:
                if 'フライ' in l_next.replace('フライドポテト', ''):
                    need_jelly = True
            if need_jelly and 'ゼリー' not in snack(ds_today):
                day_ngs[ds_today].append(
                    f'● 翌日({ds_next})ハンバーグ/フライだが前日おやつにゼリーなし'
                )

    # ── 季節チェック（クリームシチュー） ─────────────────────
    if check_seasonal_stew and 5 <= month_num <= 9:
        for ds in sorted_dates:
            if 'クリームシチュー' in lunch(ds):
                day_ngs[ds].append('● 5〜9月クリームシチューNG')

    # ── 料理名から推定される必須材料チェック ─────────────────
    # (料理名に含む文字列, 材料欄に必要なキーワードのいずれか)
    if check_name_ingredient:
        RECIPE_RULES = [
            ('おかか',   ['かつお節', 'おかか']),
            ('ごま和え', ['白ごま', '黒ごま', 'すりごま', 'ごま', '胡麻', '白胡麻', '黒胡麻', 'すり胡麻']),
            ('あんかけ', ['片栗粉']),
            ('から揚げ', ['片栗粉']),
            ('唐揚げ',   ['片栗粉']),
            ('照り焼き', ['みりん', '醤油']),
            ('蒸しパン', ['ベーキングパウダー', 'BP', 'B.P', '重曹']),
            ('味噌汁',           ['みそ', '味噌']),
            ('みそ汁',           ['みそ', '味噌']),
            ('わかめ',           ['わかめ']),  # 「わかめスープ」「わかめご飯」等
            ('果物',             FRUIT_KW),
            ('フルーツヨーグルト', FRUIT_KW),
            ('フルーツポンチ', FRUIT_KW),
        ]
        for ds in sorted_dates:
            ls_text = lunch(ds) + ' ' + snack(ds)
            i_text  = ing(ds)
            for kw, required_any in RECIPE_RULES:
                if kw in ls_text:
                    if not any(req in i_text for req in required_any):
                        missing = '・'.join(required_any)
                        day_ngs[ds].append(f'● 「{kw}」があるが材料に{missing}なし')

        # ── お菓子：市販品をそのまま提供するだけの園では材料欄に書かない運用がある
        #     （富喜屋提供の離乳食テンプレート等で確認）。ヨーグルトと同じ考え方で、
        #     園内の運用（leave-one-out）を見て判断する ─────
        bare_item_forward_check('お菓子', 'お菓子', sorted_dates, lunch, snack, ing, day_ngs)

        # ── ヨーグルト：複合料理名（フルーツヨーグルト・桃ヨーグルト等）は基本的に
        #     材料欄にも記載されるが、「飲むヨーグルト」「◯◯ヨーグルト（商品名）」の
        #     ように購入品をそのまま提供するだけの園では複合名でも材料欄に書かない
        #     運用がある（エリザベスサンダースホーム様で確認）。単体提供と同様、
        #     ファイル内の運用パターン（leave-one-out）で判定する。
        def _yogurt_compound_items(text):
            items = [re.sub(r'[（(].*?[）)]', '', i).strip() for i in text.split('/')]
            return [i for i in items if 'ヨーグルト' in i and i != 'ヨーグルト']

        _yogurt_compound_days = [
            ds for ds in sorted_dates
            if _yogurt_compound_items(lunch(ds) + '/' + snack(ds))
        ]
        _yogurt_compound_named = {ds for ds in _yogurt_compound_days if 'ヨーグルト' in ing(ds)}
        for ds in _yogurt_compound_days:
            if ds in _yogurt_compound_named:
                continue
            others = [d for d in _yogurt_compound_days if d != ds]
            if len(others) < 2:
                continue
            if sum(d in _yogurt_compound_named for d in others) / len(others) >= 0.85:
                day_ngs[ds].append('● 「ヨーグルト」があるが材料にヨーグルトなし')
        bare_item_forward_check('ヨーグルト', 'ヨーグルト', sorted_dates, lunch, snack, ing, day_ngs)

        # ── 献立名の特定食材と材料が一致しないチェック ────────
        # （例：「パインパンケーキ」なのに材料が「みかん缶」など、別の食材に
        #   すり替わっているケースを検出。表記ゆれ（パイン/パイナップル、
        #   チンゲン菜/青梗菜等）は同一視する。旧AI自由記述チェックの後継
        #   ＝ここに載っている食材については、この決定的照合だけで完結する）
        for ds in sorted_dates:
            ls_text = lunch(ds) + ' ' + snack(ds)
            ls_check = ls_text
            for idiom in FOOD_NAME_IDIOMS:
                ls_check = ls_check.replace(idiom, '')
            i_text  = ing(ds)
            for group in FOOD_SYNONYM_GROUPS:
                if any(g in ls_check for g in group) and not any(g in i_text for g in group):
                    # 「ゼリー」等の総称デザート名は、風味のもとになる果物が材料欄に
                    # なくても例外とする運用がある（海野さんチェックのルール文面にも
                    # 「ゼリーは該当食材がなくてOK」と明記）
                    if any(g in FRUIT_KW for g in group) and any(dn in ls_text for dn in FRUIT_OPTIONAL_DESSERT_KW):
                        continue
                    day_ngs[ds].append(f'● 献立名に「{group[0]}」があるが材料に見当たらない（表記が違っていないか要確認）')

        # ── 材料に果物があるのに献立名に記載がない（逆方向の表記漏れ）─────
        # この園が普段「献立名に果物を明記する」運用かをファイル内で判定し、
        # その運用が支配的な場合にだけ外れ値（記載漏れ）として報告する。
        # （果物を普段から献立名に書かない園まで一律チェックすると誤検知だらけになるため）
        # 「フルーツヨーグルト」等、複数の果物をまとめた総称名（RECIPE_RULESで
        # FRUIT_KWをそのまま必須材料とする料理名）や、ゼリー・プリン等の
        # 果物有無が日替わりの総称デザート名の日は、個別の果物名をあえて
        # 書かない運用が正当なため、母集団・判定対象から除外する。
        _fruit_generic_names = [kw for kw, req in RECIPE_RULES if req is FRUIT_KW] + FRUIT_OPTIONAL_DESSERT_KW
        _fruit_reverse_dates = [
            ds for ds in sorted_dates
            if not any(name in (lunch(ds) + ' ' + snack(ds)) for name in _fruit_generic_names)
        ]
        reverse_naming_check(FRUIT_KW, '果物', _fruit_reverse_dates, lunch, snack, ing, day_ngs)

        # ── おすまし・おすいものに「みそ」あり ─────────────────────
        for ds in sorted_dates:
            ls_text = lunch(ds) + ' ' + snack(ds)
            if ('おすまし' in ls_text or 'おすいもの' in ls_text) and ('みそ' in ing(ds) or '味噌' in ing(ds)):
                day_ngs[ds].append('● 「おすまし/おすいもの」があるが材料に「みそ」あり（不要）')

        # ── 料理名に含まれない魚が材料にあるチェック ────────────────
        # 「白身魚」は総称なので除外（白身魚のフライ→材料にタラ等があっても正常）
        _specific_fish = [f for f in FISH_KW if f != '白身魚']
        for ds in sorted_dates:
            ls_text = lunch(ds) + ' ' + snack(ds)
            i_text  = ing(ds)
            dish_fish = [f for f in _specific_fish if f in ls_text]
            ing_fish  = [f for f in _specific_fish if f in i_text]
            if dish_fish:
                for f in ing_fish:
                    if f not in dish_fish:
                        day_ngs[ds].append(f'● 材料に「{f}」があるが献立名に対応する料理なし（不要食材？）')

    # ── くにみ子ども園固有ルール ────────────────────────────────
    if check_kunimi:
        _dow_map_k = {'月': 0, '火': 1, '水': 2, '木': 3, '金': 4, '土': 5, '日': 6}

        def _week_key_k(ds):
            m2 = re.search(r'/(\d+)\(([月火水木金土日])\)', ds)
            return int(m2.group(1)) - _dow_map_k.get(m2.group(2), 0) if m2 else -1

        weeks_k = {}
        for ds in sorted_dates:
            weeks_k.setdefault(_week_key_k(ds), []).append(ds)

        # 朝おやつ：みかん缶・バナナが週1回（1〜3月はお菓子のみで果物なし）
        FRUIT_ROTATION = ['みかん缶', 'バナナ']
        if 1 <= month_num <= 3:
            for ds in sorted_dates:
                if any(f in snack_am(ds) for f in FRUIT_ROTATION):
                    day_ngs[ds].append('● 1〜3月の朝おやつは「お菓子」のみのはずが果物（みかん缶/バナナ）あり')
        else:
            for _, wds in sorted(weeks_k.items()):
                wd_k = sorted(wds, key=lambda d: _parse_date(d, year) or datetime.date.max)
                last_k = wd_k[-1]
                cnt = sum(1 for ds in wd_k for f in FRUIT_ROTATION if f in snack_am(ds))
                if cnt == 0:
                    day_ngs[last_k].append('● 今週、朝おやつに「みかん缶」「バナナ」のどちらもなし（週1回必要）')
                elif cnt >= 2:
                    day_ngs[last_k].append('● 今週、朝おやつに「みかん缶」「バナナ」が2回以上（週1回のみ）')

        # 毎月1回は入れる（麩ラスク＝麩、3時のおやつの果物、ヨーグルト和え、納豆、メロン）
        # ※麩ラスク等の完成品名は生データに出てこないため、代表食材での存在チェックで代用
        if sorted_dates:
            last_day = sorted_dates[-1]
            MONTHLY_REQUIRED = [
                ('麩ラスク（材料「麩」で代替判定）', lambda ds: snack_pm(ds), ['麩']),
                ('3時のおやつの果物',                 lambda ds: snack_pm(ds), FRUIT_KW),
                ('ヨーグルト和え（昼食に「ヨーグルト」）', lambda ds: lunch(ds), ['ヨーグルト']),
                ('納豆',                               lambda ds: ing(ds),     ['納豆']),
                ('メロン',                             lambda ds: ing(ds),     ['メロン']),
            ]
            for label, getter, kws in MONTHLY_REQUIRED:
                if not any(has(getter(ds), kws) for ds in sorted_dates):
                    day_ngs[last_day].append(f'● 今月「{label}」が一度も入っていません（毎月必須）')

            has_champon  = any('ちゃんぽん' in lunch(ds) for ds in sorted_dates)
            has_saraudon = any('皿うどん' in lunch(ds) for ds in sorted_dates)
            if not has_champon:
                day_ngs[last_day].append('● 今月「ちゃんぽん」が一度も入っていません')
            if not has_saraudon:
                day_ngs[last_day].append('● 今月「皿うどん」が一度も入っていません')

        # ビーフン・ブリは使用しない
        for ds in sorted_dates:
            whole = lunch(ds) + ' ' + snack(ds) + ' ' + ing(ds)
            if 'ビーフン' in whole:
                day_ngs[ds].append('● 「ビーフン」使用（くにみ子ども園では不使用）')
            if 'ブリ' in whole:
                day_ngs[ds].append('● 「ブリ」使用（くにみ子ども園では不使用）')

        # パン曜日ルール：火曜日はロールパン、第2・4金曜日はベーグル（パンが出る日のみ判定）
        _BREAD_KW = ['ロールパン', 'ベーグル', '食パン']
        for ds in sorted_dates:
            d = _parse_date(ds, year)
            if not d:
                continue
            l_text = lunch(ds)
            if not has(l_text, _BREAD_KW):
                continue
            dow_i = _dow(ds)
            if dow_i == 1 and 'ロールパン' not in l_text:
                day_ngs[ds].append('● 火曜日のパンは「ロールパン」のはずが違う種類')
            week_of_month = (d.day - 1) // 7 + 1
            if dow_i == 4 and week_of_month in (2, 4) and 'ベーグル' not in l_text:
                day_ngs[ds].append('● 第2・4金曜日のパンは「ベーグル」のはずが違う種類')

        # 魚とパンの組み合わせNG
        _fish_all_k = FISH_KW  # 白身魚も含めて判定
        for ds in sorted_dates:
            l_text = lunch(ds)
            if has(l_text, _BREAD_KW) and has(l_text, _fish_all_k):
                day_ngs[ds].append('● 「パン」と「魚」の組み合わせ（併用不可）')

        # ちゃんぽん・皿うどんの日：キャベツ/白菜NG、丸天+のべ板必須
        for ds in sorted_dates:
            l_text = lunch(ds)
            if 'ちゃんぽん' in l_text or '皿うどん' in l_text:
                i_text = ing(ds)
                if 'キャベツ' in i_text:
                    day_ngs[ds].append('● ちゃんぽん/皿うどんの日に「キャベツ」使用（不使用のはず）')
                if '白菜' in i_text:
                    day_ngs[ds].append('● ちゃんぽん/皿うどんの日に「白菜」使用（不使用のはず）')
                if '丸天' not in i_text:
                    day_ngs[ds].append('● ちゃんぽん/皿うどんの日に「丸天」なし（必須）')
                if 'のべ板' not in i_text:
                    day_ngs[ds].append('● ちゃんぽん/皿うどんの日に「のべ板」なし（必須）')

        # 麺の日に「ソテー」「お浸し」系の副菜が2種類（1種類にする）
        for ds in sorted_dates:
            l_text = lunch(ds)
            if has(l_text, NOODLE_KW):
                style_cnt = l_text.count('ソテー') + l_text.count('お浸し')
                if style_cnt >= 2:
                    day_ngs[ds].append('● 麺の日に「ソテー」「お浸し」系の副菜が2種類（1種類にする）')

    # ── 出力テキスト生成 ──────────────────────────────────────
    lines = [
        '【Python確定NGリスト】',
        '（AIはこの結果をそのまま採用。自分で再計算・再判断しないこと）',
        '',
    ]
    has_any = False
    for ds in sorted_dates:
        for ng in day_ngs[ds]:
            lines.append(f'{ds}: {ng}')
            has_any = True
    if not has_any:
        lines.append('（ルールベースのNG検出なし）')

    return '\n'.join(lines), day_ngs


def build_result_table(sorted_dates, entries, day_ngs):
    """day_ngs（Python確定NG）から最終結果テーブルをPythonで確定的に組み立てる。"""
    lines = ['| 日付 | 献立名 | おやつ | 結果 |', '|------|--------|--------|------|']
    for ds in sorted_dates:
        lunch_text = entries[ds]['lunch']
        snack_text = entries[ds]['snack']
        ngs = day_ngs.get(ds, [])
        result = ' ／ '.join(ngs) if ngs else 'OK'
        lines.append(f'| {ds} | {lunch_text} | {snack_text} | {result} |')
    return '\n'.join(lines)


def combine_excel_texts(texts):
    """複数シート分のexcel_text（1つの月が複数シートへ分かれている形式向け）を
    チェック実行前に1本のテキストへ結合する。
    週次・月次チェック（今月◯◯必須、週1回必須 等）はcompute_all_python_ngsが
    sorted_dates全体を見て判定するため、シートごとに個別実行して後から結果表だけ
    結合すると、月の折り返し地点で「まだ月内に1回も出てきていない」という
    誤判定（本来は他方のシート側で既に登場している）が起きる。
    必ずこの関数で結合したexcel_textを使って run_check を1回だけ呼ぶこと。"""
    return '\n'.join(t for t in texts if t.strip())


def run_check(excel_text, rules_text, file_name, sheet_name, leftover_words=None):
    """献立名⇔材料の照合を含む全チェックはPython側（compute_all_python_ngs）で決定的に行う。
    以前はここでAIに自由記述で「材料にない」と判断させ事後検証していたが、AIの言い回し次第で
    検証がすり抜ける幻覚（実在する食材を「ない」と誤診断）が繰り返し発生したため廃止した。
    献立名⇔材料の食い違いはFOOD_SYNONYM_GROUPS（compute_all_python_ngs内）に決定的な
    同義語グループとして追加していく方式に一本化している。"""
    _, day_ngs = compute_all_python_ngs(excel_text, rules_text, leftover_words)
    year, month_num, sorted_dates, entries = _parse_structured(excel_text)
    if not sorted_dates:
        return "| 日付 | 献立名 | おやつ | 結果 |\n|------|--------|--------|------|"

    return build_result_table(sorted_dates, entries, day_ngs)


with st.sidebar:
    st.title("🍱 献立チェックシステム")
    st.markdown("---")
    page = st.radio(
        "ページを選択",
        ["📋 献立チェック", "⚙️ ルール管理", "🎨 色付けルール管理", "🔄 置換ルール管理"],
        label_visibility="collapsed",
    )


if page == "📋 献立チェック":
    st.title("📋 献立チェック")

    with st.expander("🔄 0. 置換Excel生成（任意・材料名の言い換え）", expanded=False):
        st.caption(
            "「かまぼこ→のべ板」のような園独自の言い換えを先に適用したExcelを作成します。"
            "生成後にダウンロードし、その置換済みファイルを下の「1. ファイルをアップロード」で"
            "アップロードしてチェックしてください。"
        )
        all_replace_rules = load_replace_rules_list()
        if not all_replace_rules:
            st.info("置換ルールが登録されていません。「置換ルール管理」ページで追加してください。")
        else:
            replace_rule_names = [r["name"] for r in all_replace_rules]
            chosen_replace_name = st.selectbox(
                "置換ルール", replace_rule_names, key="replace_rule_selector_gen",
            )
            replace_upload = st.file_uploader(
                "置換対象のExcelファイル（.xls または .xlsx）",
                type=["xls", "xlsx"], key="replace_uploader",
            )
            if replace_upload and st.button("🔄 置換Excel生成"):
                selected_pairs = next(
                    (r["pairs"] for r in all_replace_rules if r["name"] == chosen_replace_name), []
                )
                with st.spinner("置換処理中..."):
                    replace_upload.seek(0)
                    replaced = apply_replacements_to_excel(replace_upload, selected_pairs)
                    st.session_state["replaced_excel"] = replaced.getvalue()
                    st.session_state["replaced_fname"] = replace_upload.name.rsplit(".", 1)[0]

            if st.session_state.get("replaced_excel"):
                fname_r = st.session_state.get("replaced_fname", "result")
                st.download_button(
                    "📥 置換済みExcelをダウンロード",
                    data=st.session_state["replaced_excel"],
                    file_name=f"置換済み_{fname_r}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_replaced",
                )

    st.markdown("### 1. ファイルをアップロード")
    uploaded = st.file_uploader(
        "献立Excelファイル（.xls または .xlsx）",
        type=["xls", "xlsx"],
    )

    if uploaded:
        sheets = get_sheet_names(uploaded)
        if sheets:
            st.markdown("### 2. シートを選択")
            combine_all_sheets = False
            if len(sheets) > 1:
                combine_all_sheets = st.checkbox(
                    "全シートをまとめてチェックする（月が複数シートに分かれている形式向け）",
                    key="combine_all_sheets",
                )
            selected_sheet = st.selectbox(
                "シート", sheets, disabled=combine_all_sheets,
            )

            st.markdown("### 3. 使用するルールを選択")
            all_rules = load_rules_list()
            if not all_rules:
                st.warning("ルールが登録されていません。「ルール管理」ページでルールを追加してください。")
                chosen_name = ""
                selected_rule_text = ""
            else:
                rule_names = [r["name"] for r in all_rules]
                chosen_name = st.selectbox(
                    "ルール", rule_names,
                    key="rule_selector",
                    label_visibility="collapsed"
                )
                selected_rule_text = next(
                    (r["text"] for r in all_rules if r["name"] == chosen_name), ""
                )

            all_replace_rules_check = load_replace_rules_list()
            leftover_words = []
            if all_replace_rules_check:
                replace_check_names = ["（使用しない）"] + [r["name"] for r in all_replace_rules_check]
                chosen_replace_check_name = st.selectbox(
                    "置換ルール（置換漏れチェック用・任意）",
                    replace_check_names,
                    key="replace_rule_selector_check",
                    help="事前に置換Excel生成で使ったルールを選ぶと、置換前の単語が残っていないかもチェックします。",
                )
                if chosen_replace_check_name != "（使用しない）":
                    pairs = next(
                        (r["pairs"] for r in all_replace_rules_check
                         if r["name"] == chosen_replace_check_name), []
                    )
                    leftover_words = [p["from"] for p in pairs if p.get("from")]

            st.markdown("### 4. 色付きExcel生成（目検用）")
            all_color_rules = load_color_rules_list()
            color_rule_names = [r["name"] for r in all_color_rules]
            chosen_color_name = st.selectbox(
                "色付けルール", color_rule_names,
                key="color_rule_selector",
            )
            selected_color_groups = next(
                (r["groups"] for r in all_color_rules if r["name"] == chosen_color_name),
                DEFAULT_COLOR_GROUPS,
            )
            if st.button("🎨 色付きExcel生成（全シート）"):
                with st.spinner("色付き処理中..."):
                    uploaded.seek(0)
                    colored = create_colored_excel(uploaded, selected_color_groups)
                    st.session_state["colored_excel"] = colored.getvalue()
                    st.session_state["colored_fname"] = uploaded.name.rsplit(".", 1)[0]

            if st.session_state.get("colored_excel"):
                fname_c = st.session_state.get("colored_fname", "result")
                st.download_button(
                    "📥 色付きExcelをダウンロード",
                    data=st.session_state["colored_excel"],
                    file_name=f"色付き_{fname_c}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_colored",
                )

            st.markdown("### 5. チェック開始")
            if st.button("✅ チェックを開始する", type="primary"):
                rules = selected_rule_text
                if not rules.strip():
                    st.error("ルールが選択されていないか空です。「ルール管理」ページでルールを設定してください。")
                else:
                    target_sheets = sheets if combine_all_sheets else [selected_sheet]
                    with st.spinner("チェック中です..."):
                        try:
                            excel_texts = []
                            for sn in target_sheets:
                                uploaded.seek(0)
                                excel_texts.append(excel_to_text(uploaded, sn))
                            # 週次・月次チェックが月全体を通して正しく判定されるよう、
                            # シートごとに個別実行せず、結合後のテキストで1回だけ実行する
                            combined_text = combine_excel_texts(excel_texts)
                            result = run_check(
                                combined_text, rules, uploaded.name,
                                '/'.join(target_sheets), leftover_words,
                            )
                            st.session_state["last_result"] = result
                            st.session_state["last_filename"] = uploaded.name
                            st.session_state["last_sheet"] = (
                                f"全シート（{', '.join(target_sheets)}）"
                                if combine_all_sheets else selected_sheet
                            )
                            st.session_state["last_rule_name"] = chosen_name
                        except Exception as e:
                            st.error(f"エラーが発生しました: {e}")
                            result = None

            if st.session_state.get("last_result"):
                st.markdown("---")
                _fname_disp = st.session_state.get("last_filename", "")
                _sheet_disp = st.session_state.get("last_sheet", "")
                _rule_disp = st.session_state.get("last_rule_name", "")
                st.subheader(f"チェック結果 ― {_fname_disp}　シート: {_sheet_disp}")
                if _rule_disp:
                    st.caption(f"使用ルール：{_rule_disp}")
                result_text = st.session_state["last_result"]
                st.markdown(result_text)
                fname = st.session_state.get("last_filename", "result").rsplit(".", 1)[0]
                docx_data = table_to_docx(result_text, title=f"チェック結果 ― {_fname_disp}　シート: {_sheet_disp}")
                st.download_button(
                    "📥 結果をWordファイルで保存",
                    data=docx_data,
                    file_name=f"チェック結果_{fname}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_word",
                )


elif page == "⚙️ ルール管理":
    st.title("⚙️ ルール管理")
    st.caption("ルールを登録・編集すると、チェック画面のセレクトボックスに表示されます。保存はGitHubにも自動反映されます。")

    all_rules = load_rules_list()

    # ── 編集フォーム（新規追加 or 既存編集）────────────────────
    editing = st.session_state.get("rule_editing")  # {"id": str|None, "name": str, "text": str}

    if editing is not None:
        is_new = editing["id"] is None
        st.subheader("➕ 新しいルールを追加" if is_new else f"✏️ 編集：{editing['name']}")
        with st.form("rule_form"):
            new_name = st.text_input(
                "ルール名（例：さかえ保育園、おおみや共通 など）",
                value=editing["name"],
                max_chars=50,
            )
            new_text = st.text_area(
                "ルール本文",
                value=editing["text"],
                height=500,
                help="このルールをAIに渡してチェックします。",
            )
            col_s, col_c = st.columns([1, 5])
            with col_s:
                submitted = st.form_submit_button("💾 保存", type="primary")
            with col_c:
                cancelled = st.form_submit_button("✖ キャンセル")

        if submitted:
            if not new_name.strip():
                st.error("ルール名を入力してください。")
            else:
                if is_new:
                    new_id = str(int(datetime.datetime.now().timestamp() * 1000))
                    all_rules.append({"id": new_id, "name": new_name.strip(), "text": new_text})
                else:
                    for r in all_rules:
                        if r["id"] == editing["id"]:
                            r["name"] = new_name.strip()
                            r["text"] = new_text
                            break
                save_rules_list(all_rules)
                ok, msg = push_rules_list_to_github(all_rules)
                if ok:
                    st.success(msg)
                else:
                    st.warning(f"アプリには保存済みです。GitHub更新に失敗しました：{msg}")
                del st.session_state["rule_editing"]
                st.rerun()
        if cancelled:
            del st.session_state["rule_editing"]
            st.rerun()

    else:
        if st.button("➕ 新しいルールを追加", type="primary"):
            st.session_state["rule_editing"] = {"id": None, "name": "", "text": ""}
            st.rerun()

    # ── ルール一覧 ────────────────────────────────────────────
    if editing is None:
        st.markdown("---")
        if not all_rules:
            st.info("ルールがまだ登録されていません。上のボタンから追加してください。")
        else:
            st.markdown(f"**登録済みルール：{len(all_rules)} 件**")
            for rule in all_rules:
                with st.container(border=True):
                    col_n, col_e, col_d = st.columns([6, 1, 1])
                    with col_n:
                        st.markdown(f"**{rule['name']}**")
                        preview = rule["text"][:80].replace("\n", "  ") + ("…" if len(rule["text"]) > 80 else "")
                        st.caption(preview)
                    with col_e:
                        if st.button("✏ 編集", key=f"edit_{rule['id']}"):
                            st.session_state["rule_editing"] = {
                                "id": rule["id"],
                                "name": rule["name"],
                                "text": rule["text"],
                            }
                            st.rerun()
                    with col_d:
                        # 削除: 1回目クリックで確認フラグ、2回目で実行
                        confirm_key = f"confirm_del_{rule['id']}"
                        if st.session_state.get(confirm_key):
                            if st.button("本当に削除", key=f"yes_{rule['id']}",
                                         type="primary"):
                                all_rules = [r for r in all_rules if r["id"] != rule["id"]]
                                save_rules_list(all_rules)
                                push_rules_list_to_github(all_rules)
                                del st.session_state[confirm_key]
                                st.rerun()
                        else:
                            if st.button("🗑 削除", key=f"del_{rule['id']}"):
                                st.session_state[confirm_key] = True
                                st.rerun()


elif page == "🎨 色付けルール管理":
    st.title("🎨 色付けルール管理")
    st.caption(
        "「どのキーワードにどの色を付けるか」をグループ単位で登録できます。"
        "チェック画面の「色付きExcel生成」でルールを選んで使えます。保存はGitHubにも自動反映されます。"
    )

    all_color_rules = load_color_rules_list()
    editing = st.session_state.get("color_rule_editing")
    # editing = {"id": str|None, "name": str, "groups": [{"keywords": [str,...], "color": "RRGGBB"}, ...]}

    if editing is not None:
        is_new = editing["id"] is None
        st.subheader("➕ 新しい色付けルールを追加" if is_new else f"✏️ 編集：{editing['name']}")

        new_name = st.text_input(
            "ルール名（例：さかえ保育園向け配色 など）",
            value=editing["name"],
            max_chars=50,
            key="color_rule_name_input",
        )

        st.markdown("**色グループ**（同じ色を付けたいキーワードをカンマ区切りでまとめて入力）")
        groups = editing["groups"]
        ver = st.session_state.setdefault("color_rule_edit_version", 0)

        for i, g in enumerate(groups):
            col_kw, col_color, col_del = st.columns([5, 1, 1])
            with col_kw:
                kw_text = ", ".join(g["keywords"])
                new_kw = st.text_input(
                    f"キーワード群 {i + 1}", value=kw_text,
                    key=f"color_group_kw_{ver}_{i}", label_visibility="collapsed",
                    placeholder="例：コーン, 人参, かぼちゃ",
                )
                g["keywords"] = [k.strip() for k in new_kw.split(",") if k.strip()]
            with col_color:
                new_color = st.color_picker(
                    f"色 {i + 1}", value=f"#{g['color']}",
                    key=f"color_group_color_{ver}_{i}", label_visibility="collapsed",
                )
                g["color"] = new_color.lstrip("#").upper()
            with col_del:
                if st.button("🗑", key=f"color_group_del_{ver}_{i}"):
                    groups.pop(i)
                    st.session_state["color_rule_edit_version"] = ver + 1
                    st.rerun()

        if st.button("＋ グループを追加", key="color_group_add"):
            groups.append({"keywords": [], "color": "FFFFFF"})
            st.session_state["color_rule_edit_version"] = ver + 1
            st.rerun()

        st.markdown("---")
        col_s, col_c = st.columns([1, 5])
        with col_s:
            save_clicked = st.button("💾 保存", type="primary", key="color_rule_save")
        with col_c:
            cancel_clicked = st.button("✖ キャンセル", key="color_rule_cancel")

        if save_clicked:
            if not new_name.strip():
                st.error("ルール名を入力してください。")
            else:
                clean_groups = [g for g in groups if g["keywords"]]
                if is_new:
                    new_id = str(int(datetime.datetime.now().timestamp() * 1000))
                    all_color_rules.append(
                        {"id": new_id, "name": new_name.strip(), "groups": clean_groups}
                    )
                else:
                    found = False
                    for r in all_color_rules:
                        if r["id"] == editing["id"]:
                            r["name"] = new_name.strip()
                            r["groups"] = clean_groups
                            found = True
                            break
                    if not found:
                        all_color_rules.append(
                            {"id": editing["id"], "name": new_name.strip(), "groups": clean_groups}
                        )
                save_color_rules_list(all_color_rules)
                ok, msg = push_color_rules_list_to_github(all_color_rules)
                if ok:
                    st.success(msg)
                else:
                    st.warning(f"アプリには保存済みです。GitHub更新に失敗しました：{msg}")
                del st.session_state["color_rule_editing"]
                st.session_state.pop("color_rule_edit_version", None)
                st.rerun()
        if cancel_clicked:
            del st.session_state["color_rule_editing"]
            st.session_state.pop("color_rule_edit_version", None)
            st.rerun()

    else:
        if st.button("➕ 新しい色付けルールを追加", type="primary", key="color_rule_new"):
            st.session_state["color_rule_editing"] = {"id": None, "name": "", "groups": []}
            st.rerun()

        st.markdown("---")
        if not all_color_rules:
            st.info("色付けルールがまだ登録されていません。上のボタンから追加してください。")
        else:
            st.markdown(f"**登録済み色付けルール：{len(all_color_rules)} 件**")
            for rule in all_color_rules:
                with st.container(border=True):
                    col_n, col_e, col_d = st.columns([6, 1, 1])
                    with col_n:
                        st.markdown(f"**{rule['name']}**")
                        for g in rule["groups"]:
                            swatch = (
                                f"<span style='display:inline-block;width:14px;height:14px;"
                                f"background:#{g['color']};border:1px solid #999;"
                                f"margin-right:6px;vertical-align:middle;'></span>"
                            )
                            st.markdown(f"{swatch}{', '.join(g['keywords'])}", unsafe_allow_html=True)
                    with col_e:
                        if st.button("✏ 編集", key=f"color_edit_{rule['id']}"):
                            st.session_state["color_rule_editing"] = {
                                "id": rule["id"],
                                "name": rule["name"],
                                "groups": [dict(g) for g in rule["groups"]],
                            }
                            st.rerun()
                    with col_d:
                        confirm_key = f"color_confirm_del_{rule['id']}"
                        if st.session_state.get(confirm_key):
                            if st.button("本当に削除", key=f"color_yes_{rule['id']}", type="primary"):
                                all_color_rules = [r for r in all_color_rules if r["id"] != rule["id"]]
                                save_color_rules_list(all_color_rules)
                                push_color_rules_list_to_github(all_color_rules)
                                del st.session_state[confirm_key]
                                st.rerun()
                        else:
                            if st.button("🗑 削除", key=f"color_del_{rule['id']}"):
                                st.session_state[confirm_key] = True
                                st.rerun()


elif page == "🔄 置換ルール管理":
    st.title("🔄 置換ルール管理")
    st.caption(
        "「材料名A→材料名B」のような園独自の言い換えペアを複数登録できます。"
        "チェック画面の「置換Excel生成」で使うほか、チェック実行時に「置換前の単語」が"
        "残っていないかの検出にも使えます。保存はGitHubにも自動反映されます。"
    )

    all_replace_rules = load_replace_rules_list()
    editing = st.session_state.get("replace_rule_editing")
    # editing = {"id": str|None, "name": str, "pairs": [{"from": str, "to": str}, ...]}

    if editing is not None:
        is_new = editing["id"] is None
        st.subheader("➕ 新しい置換ルールを追加" if is_new else f"✏️ 編集：{editing['name']}")

        new_name = st.text_input(
            "ルール名（例：くにみ子ども園 置換ルール など）",
            value=editing["name"],
            max_chars=50,
            key="replace_rule_name_input",
        )

        st.markdown("**置換ペア**（置換前 → 置換後）")
        pairs = editing["pairs"]
        ver = st.session_state.setdefault("replace_rule_edit_version", 0)

        for i, p in enumerate(pairs):
            col_from, col_to, col_del = st.columns([3, 3, 1])
            with col_from:
                new_from = st.text_input(
                    f"置換前 {i + 1}", value=p["from"],
                    key=f"replace_pair_from_{ver}_{i}", label_visibility="collapsed",
                    placeholder="例：かまぼこ",
                )
                p["from"] = new_from.strip()
            with col_to:
                new_to = st.text_input(
                    f"置換後 {i + 1}", value=p["to"],
                    key=f"replace_pair_to_{ver}_{i}", label_visibility="collapsed",
                    placeholder="例：のべ板",
                )
                p["to"] = new_to.strip()
            with col_del:
                if st.button("🗑", key=f"replace_pair_del_{ver}_{i}"):
                    pairs.pop(i)
                    st.session_state["replace_rule_edit_version"] = ver + 1
                    st.rerun()

        if st.button("＋ ペアを追加", key="replace_pair_add"):
            pairs.append({"from": "", "to": ""})
            st.session_state["replace_rule_edit_version"] = ver + 1
            st.rerun()

        st.markdown("---")
        col_s, col_c = st.columns([1, 5])
        with col_s:
            save_clicked = st.button("💾 保存", type="primary", key="replace_rule_save")
        with col_c:
            cancel_clicked = st.button("✖ キャンセル", key="replace_rule_cancel")

        if save_clicked:
            if not new_name.strip():
                st.error("ルール名を入力してください。")
            else:
                clean_pairs = [p for p in pairs if p["from"]]
                if is_new:
                    new_id = str(int(datetime.datetime.now().timestamp() * 1000))
                    all_replace_rules.append(
                        {"id": new_id, "name": new_name.strip(), "pairs": clean_pairs}
                    )
                else:
                    found = False
                    for r in all_replace_rules:
                        if r["id"] == editing["id"]:
                            r["name"] = new_name.strip()
                            r["pairs"] = clean_pairs
                            found = True
                            break
                    if not found:
                        all_replace_rules.append(
                            {"id": editing["id"], "name": new_name.strip(), "pairs": clean_pairs}
                        )
                save_replace_rules_list(all_replace_rules)
                ok, msg = push_replace_rules_list_to_github(all_replace_rules)
                if ok:
                    st.success(msg)
                else:
                    st.warning(f"アプリには保存済みです。GitHub更新に失敗しました：{msg}")
                del st.session_state["replace_rule_editing"]
                st.session_state.pop("replace_rule_edit_version", None)
                st.rerun()
        if cancel_clicked:
            del st.session_state["replace_rule_editing"]
            st.session_state.pop("replace_rule_edit_version", None)
            st.rerun()

    else:
        if st.button("➕ 新しい置換ルールを追加", type="primary", key="replace_rule_new"):
            st.session_state["replace_rule_editing"] = {"id": None, "name": "", "pairs": []}
            st.rerun()

        st.markdown("---")
        if not all_replace_rules:
            st.info("置換ルールがまだ登録されていません。上のボタンから追加してください。")
        else:
            st.markdown(f"**登録済み置換ルール：{len(all_replace_rules)} 件**")
            for rule in all_replace_rules:
                with st.container(border=True):
                    col_n, col_e, col_d = st.columns([6, 1, 1])
                    with col_n:
                        st.markdown(f"**{rule['name']}**")
                        preview = " ／ ".join(f"{p['from']}→{p['to']}" for p in rule["pairs"][:6])
                        if len(rule["pairs"]) > 6:
                            preview += " …"
                        st.caption(preview or "（ペア未登録）")
                    with col_e:
                        if st.button("✏ 編集", key=f"replace_edit_{rule['id']}"):
                            st.session_state["replace_rule_editing"] = {
                                "id": rule["id"],
                                "name": rule["name"],
                                "pairs": [dict(p) for p in rule["pairs"]],
                            }
                            st.rerun()
                    with col_d:
                        confirm_key = f"replace_confirm_del_{rule['id']}"
                        if st.session_state.get(confirm_key):
                            if st.button("本当に削除", key=f"replace_yes_{rule['id']}", type="primary"):
                                all_replace_rules = [r for r in all_replace_rules if r["id"] != rule["id"]]
                                save_replace_rules_list(all_replace_rules)
                                push_replace_rules_list_to_github(all_replace_rules)
                                del st.session_state[confirm_key]
                                st.rerun()
                        else:
                            if st.button("🗑 削除", key=f"replace_del_{rule['id']}"):
                                st.session_state[confirm_key] = True
                                st.rerun()
